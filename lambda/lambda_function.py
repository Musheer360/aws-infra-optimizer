import json
import boto3
import os
import csv
from datetime import datetime, timedelta, timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
from io import BytesIO, StringIO
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# Pure enrichment helpers (platform->pricing map, Compute Optimizer option
# selection, gp2/gp3 sizing, effort/risk/priority scoring). Kept in a sibling
# module so the accuracy-critical logic is unit-testable. Import defensively so
# a packaging hiccup never takes down the whole scan.
try:
    import enrichment
except ImportError:  # pragma: no cover - fallback for unusual import paths
    from . import enrichment  # type: ignore

# Template path - CloudThat letterhead template
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'template.docx')

# Global pricing cache (persists across Lambda invocations)
PRICING_CACHE = {}
CACHE_TTL = 3600  # 1 hour in seconds

# Comprehensive region to location mapping for AWS Pricing API
REGION_LOCATION_MAP = {
    'us-east-1': 'US East (N. Virginia)',
    'us-east-2': 'US East (Ohio)',
    'us-west-1': 'US West (N. California)',
    'us-west-2': 'US West (Oregon)',
    'af-south-1': 'Africa (Cape Town)',
    'ap-east-1': 'Asia Pacific (Hong Kong)',
    'ap-south-1': 'Asia Pacific (Mumbai)',
    'ap-south-2': 'Asia Pacific (Hyderabad)',
    'ap-southeast-1': 'Asia Pacific (Singapore)',
    'ap-southeast-2': 'Asia Pacific (Sydney)',
    'ap-southeast-3': 'Asia Pacific (Jakarta)',
    'ap-southeast-4': 'Asia Pacific (Melbourne)',
    'ap-northeast-1': 'Asia Pacific (Tokyo)',
    'ap-northeast-2': 'Asia Pacific (Seoul)',
    'ap-northeast-3': 'Asia Pacific (Osaka)',
    'ca-central-1': 'Canada (Central)',
    'ca-west-1': 'Canada West (Calgary)',
    'eu-central-1': 'EU (Frankfurt)',
    'eu-central-2': 'Europe (Zurich)',
    'eu-west-1': 'EU (Ireland)',
    'eu-west-2': 'EU (London)',
    'eu-west-3': 'EU (Paris)',
    'eu-south-1': 'EU (Milan)',
    'eu-south-2': 'Europe (Spain)',
    'eu-north-1': 'EU (Stockholm)',
    'il-central-1': 'Israel (Tel Aviv)',
    'me-south-1': 'Middle East (Bahrain)',
    'me-central-1': 'Middle East (UAE)',
    'sa-east-1': 'South America (Sao Paulo)',
}

def lambda_handler(event, context):
    try:
        body = json.loads(event.get('body', '{}'))
    except (json.JSONDecodeError, TypeError):
        return {
            'statusCode': 400,
            'headers': {
                'Content-Type': 'application/json',
            },
            'body': json.dumps({'message': 'Invalid or missing request body'})
        }
    
    client_name = body.get('clientName', 'Client')
    export_format = body.get('exportFormat', 'docx')

    try:
        result = run_full_scan(body)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json'},
            'body': json.dumps({'message': f'Scan failed: {e}'})
        }

    content, filename = make_report(result, client_name, export_format)
    summary = scan_result_summary(result)
    summary['file'] = base64.b64encode(content).decode('utf-8')
    summary['filename'] = filename

    return {
        'statusCode': 200,
        'headers': {'Content-Type': 'application/json'},
        'body': json.dumps(summary, default=str)
    }

def _get_ec2_instance_meta(ec2, instance_id):
    """Return (platform_details, tenancy, tags_dict) for an instance.

    Used so Compute Optimizer recommendations are priced with the instance's
    real OS/tenancy rather than assuming Linux/Shared.
    """
    platform_details = 'Linux/UNIX'
    tenancy = 'default'
    tags = {}
    try:
        resp = ec2.describe_instances(InstanceIds=[instance_id])
        inst = resp['Reservations'][0]['Instances'][0]
        platform_details = inst.get('PlatformDetails', 'Linux/UNIX')
        tenancy = inst.get('Placement', {}).get('Tenancy', 'default')
        tags = get_resource_tags(inst.get('Tags', []))
    except Exception:
        pass
    return platform_details, tenancy, tags


def _co_reason(rec):
    """Build a readable reason string from a Compute Optimizer recommendation."""
    finding = rec.get('finding', 'Optimizable')
    codes = rec.get('findingReasonCodes') or []
    if codes:
        pretty = ', '.join(c.replace('CPUOverprovisioned', 'CPU over-provisioned')
                            .replace('MemoryOverprovisioned', 'Memory over-provisioned')
                            .replace('EBSThroughputOverprovisioned', 'EBS throughput over-provisioned')
                            .replace('NetworkBandwidthOverprovisioned', 'Network over-provisioned')
                            for c in codes[:3])
        return f"{finding} ({pretty})"
    return finding


def scan_ec2_instances(session):
    recommendations = []
    skipped_resources = []
    ec2 = session.client('ec2')
    cloudwatch = session.client('cloudwatch')
    compute_optimizer = session.client('compute-optimizer')
    
    # Minimum data points required for reliable analysis (at least 7 days of data)
    MIN_DATA_POINTS = 7
    
    # Check for Reserved Instances
    reserved_instances = {}
    try:
        ris = ec2.describe_reserved_instances(Filters=[{'Name': 'state', 'Values': ['active']}])
        for ri in ris['ReservedInstances']:
            instance_type = ri['InstanceType']
            count = ri['InstanceCount']
            reserved_instances[instance_type] = reserved_instances.get(instance_type, 0) + count
    except Exception as e:
        print(f"RI check error: {e}")
    
    # Try Compute Optimizer first (with pagination)
    try:
        paginator_token = None
        while True:
            params = {}
            if paginator_token:
                params['nextToken'] = paginator_token
            
            response = compute_optimizer.get_ec2_instance_recommendations(**params)
            
            for rec in response.get('instanceRecommendations', []):
                if rec['finding'] in ['Overprovisioned', 'Underprovisioned']:
                    instance_id = rec['instanceArn'].split('/')[-1]
                    current_type = rec['currentInstanceType']
                    
                    # Skip if covered by Reserved Instance
                    if reserved_instances.get(current_type, 0) > 0:
                        reserved_instances[current_type] -= 1
                        continue
                    
                    # Select AWS's recommended option correctly. AWS ranks
                    # options (rank 1 = balanced top pick); the previous
                    # min(projectedUtilization) heuristic selected the LARGEST
                    # instance and understated savings.
                    best = enrichment.pick_compute_optimizer_option(rec)
                    if best and best.get('instanceType'):
                        recommended_type = best['instanceType']

                        # Instance platform/tenancy/tags -> OS-aware pricing.
                        platform_details, tenancy_raw, tags = _get_ec2_instance_meta(ec2, instance_id)
                        os_name, pre_sw = enrichment.map_platform_to_pricing(platform_details)
                        tenancy = enrichment.tenancy_to_pricing(tenancy_raw)

                        # Prefer AWS's own savings estimate (nets out RI/SP
                        # discounts when available) over recomputing it.
                        co_savings, basis = enrichment.co_option_savings(best)

                        try:
                            # Calculate savings with actual pricing (no fallback - must get real data)
                            current_cost = get_instance_cost(current_type, session.region_name,
                                                             os_name, tenancy, pre_sw)
                            recommended_cost = get_instance_cost(recommended_type, session.region_name,
                                                                os_name, tenancy, pre_sw)
                            price_savings = (current_cost - recommended_cost) * 730

                            if co_savings is not None and co_savings > 0:
                                monthly_savings = round(co_savings, 2)
                                savings_basis = basis
                            else:
                                monthly_savings = round(price_savings, 2)
                                savings_basis = 'on_demand'

                            if monthly_savings > 0:
                                recommendations.append({
                                    'instance_id': instance_id,
                                    'current_type': current_type,
                                    'recommended_type': recommended_type,
                                    'current_cost': round(current_cost * 730, 2),
                                    'recommended_cost': round(recommended_cost * 730, 2),
                                    'monthly_savings': monthly_savings,
                                    'savings_basis': savings_basis,
                                    'platform': platform_details,
                                    'source': 'Compute Optimizer',
                                    'performance_risk': best.get('performanceRisk', 'N/A'),
                                    'reason': _co_reason(rec),
                                    'confidence': 'High',
                                    'effort': 'Medium',
                                    'risk': 'Medium',
                                    'cpu_avg': get_metric_value(rec, 'CPU'),
                                    'memory_avg': get_metric_value(rec, 'MEMORY'),
                                    'tags': tags
                                })
                        except PricingUnavailableError as e:
                            skipped_resources.append(f"EC2 {instance_id}: {e}")
                            print(f"Skipping EC2 {instance_id} - pricing unavailable: {e}")
            
            paginator_token = response.get('nextToken')
            if not paginator_token:
                break
                
    except Exception as e:
        print(f"Compute Optimizer not available: {e}")
    
    # Fallback: Check for low utilization via CloudWatch (with pagination)
    try:
        # Get all running instances with pagination
        all_instances = []
        paginator = ec2.get_paginator('describe_instances')
        for page in paginator.paginate(Filters=[{'Name': 'instance-state-name', 'Values': ['running']}]):
            for reservation in page['Reservations']:
                all_instances.extend(reservation['Instances'])
        
        # Check for Auto Scaling groups (with pagination)
        asg_instances = set()
        try:
            asg_client = session.client('autoscaling')
            asg_paginator = asg_client.get_paginator('describe_auto_scaling_groups')
            for page in asg_paginator.paginate():
                for asg in page['AutoScalingGroups']:
                    for instance in asg['Instances']:
                        asg_instances.add(instance['InstanceId'])
        except Exception as e:
            print(f"ASG check warning: {e}")
        
        for instance in all_instances:
            instance_id = instance['InstanceId']
            instance_type = instance['InstanceType']
            
            # Skip if already in Compute Optimizer recommendations
            if any(r['instance_id'] == instance_id for r in recommendations):
                continue
            
            # Skip if in Auto Scaling group
            if instance_id in asg_instances:
                continue
            
            # Skip if covered by Reserved Instance
            if reserved_instances.get(instance_type, 0) > 0:
                reserved_instances[instance_type] -= 1
                continue
            
            # Check CPU utilization - use 14 days of data
            end_time = datetime.now(timezone.utc)
            start_time = end_time - timedelta(days=14)
            
            cpu_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/EC2',
                MetricName='CPUUtilization',
                Dimensions=[{'Name': 'InstanceId', 'Value': instance_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            # Also check network utilization to ensure instance isn't network-bound
            network_in_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/EC2',
                MetricName='NetworkIn',
                Dimensions=[{'Name': 'InstanceId', 'Value': instance_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            network_out_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/EC2',
                MetricName='NetworkOut',
                Dimensions=[{'Name': 'InstanceId', 'Value': instance_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            # Require minimum data points to ensure we have reliable data
            if not cpu_stats['Datapoints'] or len(cpu_stats['Datapoints']) < MIN_DATA_POINTS:
                print(f"Skipping EC2 {instance_id} - insufficient data points ({len(cpu_stats.get('Datapoints', []))} < {MIN_DATA_POINTS})")
                continue
            
            avg_cpu = sum(d['Average'] for d in cpu_stats['Datapoints']) / len(cpu_stats['Datapoints'])
            max_cpu = max(d['Maximum'] for d in cpu_stats['Datapoints'])
            
            # Check network utilization (bytes/day) - skip if high network usage
            # High network could indicate the instance is sized for network, not CPU
            max_network_in = max(d['Maximum'] for d in network_in_stats['Datapoints']) if network_in_stats['Datapoints'] else 0
            max_network_out = max(d['Maximum'] for d in network_out_stats['Datapoints']) if network_out_stats['Datapoints'] else 0
            
            # Skip if network usage is high (> 1GB/day peak) as instance may be network-bound
            if max_network_in > 1_000_000_000 or max_network_out > 1_000_000_000:
                print(f"Skipping EC2 {instance_id} - high network utilization")
                continue
            
            # CONSERVATIVE thresholds to avoid underprovisioning in production
            # Only recommend downsizing if BOTH average AND max CPU are very low
            # This ensures we have significant headroom for load spikes
            if avg_cpu < 5 and max_cpu < 15:
                try:
                    # OS-aware pricing from the instance's own platform/tenancy.
                    platform_details = instance.get('PlatformDetails', 'Linux/UNIX')
                    os_name, pre_sw = enrichment.map_platform_to_pricing(platform_details)
                    tenancy = enrichment.tenancy_to_pricing(instance.get('Placement', {}).get('Tenancy', 'default'))
                    current_cost = get_instance_cost(instance_type, session.region_name,
                                                     os_name, tenancy, pre_sw)
                    # Recommend one size smaller
                    smaller_type = get_smaller_instance_type(instance_type)
                    if smaller_type:
                        smaller_cost = get_instance_cost(smaller_type, session.region_name,
                                                         os_name, tenancy, pre_sw)
                        monthly_savings = (current_cost - smaller_cost) * 730
                        
                        if monthly_savings > 0:
                            recommendations.append({
                                'instance_id': instance_id,
                                'current_type': instance_type,
                                'recommended_type': smaller_type,
                                'current_cost': round(current_cost * 730, 2),
                                'recommended_cost': round(smaller_cost * 730, 2),
                                'monthly_savings': round(monthly_savings, 2),
                                'savings_basis': 'on_demand',
                                'platform': platform_details,
                                'source': 'CloudWatch',
                                'reason': f'Very low CPU utilization (avg: {avg_cpu:.1f}%, max: {max_cpu:.1f}%)',
                                'confidence': 'Medium',
                                'effort': 'Medium',
                                'risk': 'Medium',
                                'cpu_avg': round(avg_cpu, 1),
                                'memory_avg': 'N/A',
                                'data_points': len(cpu_stats['Datapoints']),
                                'tags': get_resource_tags(instance.get('Tags', []))
                            })
                except PricingUnavailableError as e:
                    skipped_resources.append(f"EC2 {instance_id}: {e}")
                    print(f"Skipping EC2 {instance_id} - pricing unavailable: {e}")
    except Exception as e:
        print(f"CloudWatch fallback error: {e}")
    
    return recommendations


def scan_ebs_volumes(session):
    recommendations = []
    skipped_resources = []
    ec2 = session.client('ec2')
    
    try:
        # Get all volumes with pagination
        all_volumes = []
        paginator = ec2.get_paginator('describe_volumes')
        for page in paginator.paginate():
            all_volumes.extend(page['Volumes'])
        
        for volume in all_volumes:
            volume_id = volume['VolumeId']
            volume_type = volume['VolumeType']
            size = volume['Size']
            state = volume['State']
            iops = volume.get('Iops', 0)
            throughput = volume.get('Throughput', 0)  # Only for gp3
            
            try:
                # Unattached volumes
                if state == 'available':
                    monthly_cost = calculate_ebs_cost(volume_type, size, session.region_name, iops, throughput)
                    recommendations.append({
                        'volume_id': volume_id,
                        'size': size,
                        'type': volume_type,
                        'issue': 'Unattached',
                        'recommendation': 'Delete if not needed or attach to instance',
                        'monthly_savings': round(monthly_cost, 2),
                        'confidence': 'High',
                        'tags': get_resource_tags(volume.get('Tags', []))
                    })
                
                # gp2 to gp3 migration
                elif volume_type == 'gp2':
                    current_cost = calculate_ebs_cost('gp2', size, session.region_name, iops, throughput)
                    # A faithful gp3 replacement must match gp2's baseline
                    # performance. gp2 provisions 3 IOPS/GiB, so volumes > ~334
                    # GiB already exceed gp3's free 3,000 IOPS / 125 MB/s and must
                    # pay for the extra - otherwise we overstate savings.
                    gp3_iops, gp3_throughput = enrichment.gp3_target_performance(size)
                    gp3_cost = calculate_ebs_cost('gp3', size, session.region_name, gp3_iops, gp3_throughput)
                    monthly_savings = current_cost - gp3_cost

                    if monthly_savings > 0:
                        note = ''
                        if gp3_iops > 3000 or gp3_throughput > 125:
                            note = f' (gp3 sized to {gp3_iops} IOPS / {gp3_throughput} MB/s to match gp2 baseline)'
                        recommendations.append({
                            'volume_id': volume_id,
                            'size': size,
                            'type': volume_type,
                            'issue': 'Using gp2',
                            'recommendation': f'Migrate to gp3 for better performance and lower cost{note}',
                            'monthly_savings': round(monthly_savings, 2),
                            'confidence': 'High',
                            'tags': get_resource_tags(volume.get('Tags', []))
                        })
            except PricingUnavailableError as e:
                skipped_resources.append(f"EBS {volume_id}: {e}")
                print(f"Skipping EBS {volume_id} - pricing unavailable: {e}")
    except Exception as e:
        print(f"EBS scan error: {e}")
    
    return recommendations


def scan_rds_instances(session):
    recommendations = []
    skipped_resources = []
    rds = session.client('rds')
    cloudwatch = session.client('cloudwatch')
    
    # Minimum data points required for reliable analysis (at least 7 days of data)
    MIN_DATA_POINTS = 7
    
    try:
        # Get all RDS instances with pagination
        all_instances = []
        paginator = rds.get_paginator('describe_db_instances')
        for page in paginator.paginate():
            all_instances.extend(page['DBInstances'])
        
        for db in all_instances:
            db_id = db['DBInstanceIdentifier']
            db_class = db['DBInstanceClass']
            engine = db['Engine']
            multi_az = db.get('MultiAZ', False)
            
            # Check CPU utilization - 14 days of data
            end_time = datetime.now(timezone.utc)
            start_time = end_time - timedelta(days=14)
            
            cpu_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/RDS',
                MetricName='CPUUtilization',
                Dimensions=[{'Name': 'DBInstanceIdentifier', 'Value': db_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            conn_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/RDS',
                MetricName='DatabaseConnections',
                Dimensions=[{'Name': 'DBInstanceIdentifier', 'Value': db_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            # Also check memory utilization (FreeableMemory) and I/O metrics
            memory_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/RDS',
                MetricName='FreeableMemory',
                Dimensions=[{'Name': 'DBInstanceIdentifier', 'Value': db_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Minimum']
            )
            
            read_iops_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/RDS',
                MetricName='ReadIOPS',
                Dimensions=[{'Name': 'DBInstanceIdentifier', 'Value': db_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            write_iops_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/RDS',
                MetricName='WriteIOPS',
                Dimensions=[{'Name': 'DBInstanceIdentifier', 'Value': db_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            # Require minimum data points for all critical metrics
            if not cpu_stats['Datapoints'] or len(cpu_stats['Datapoints']) < MIN_DATA_POINTS:
                print(f"Skipping RDS {db_id} - insufficient CPU data points")
                continue
            
            if not conn_stats['Datapoints'] or len(conn_stats['Datapoints']) < MIN_DATA_POINTS:
                print(f"Skipping RDS {db_id} - insufficient connection data points")
                continue
            
            avg_cpu = sum(d['Average'] for d in cpu_stats['Datapoints']) / len(cpu_stats['Datapoints'])
            max_cpu = max(d['Maximum'] for d in cpu_stats['Datapoints'])
            avg_conn = sum(d['Average'] for d in conn_stats['Datapoints']) / len(conn_stats['Datapoints'])
            max_conn = max(d['Maximum'] for d in conn_stats['Datapoints'])
            
            # Check if memory is under pressure (low freeable memory could indicate memory-bound workload)
            min_freeable_memory = min(d['Minimum'] for d in memory_stats['Datapoints']) if memory_stats['Datapoints'] else None
            
            # Check I/O utilization - high IOPS could indicate I/O bound workload
            max_read_iops = max(d['Maximum'] for d in read_iops_stats['Datapoints']) if read_iops_stats['Datapoints'] else 0
            max_write_iops = max(d['Maximum'] for d in write_iops_stats['Datapoints']) if write_iops_stats['Datapoints'] else 0
            
            # Skip if memory seems constrained. Use a *relative* threshold
            # (<10% of the class's RAM free at the trough) instead of a flat
            # 500 MB, which is meaningless across instance sizes.
            class_mem_gb = enrichment.rds_class_memory_gb(db_class)
            if min_freeable_memory is not None:
                if class_mem_gb:
                    if min_freeable_memory < 0.10 * class_mem_gb * (1024 ** 3):
                        print(f"Skipping RDS {db_id} - memory appears constrained (<10% of {class_mem_gb}GiB free)")
                        continue
                elif min_freeable_memory < 500_000_000:
                    # Fallback when the class RAM is unknown.
                    print(f"Skipping RDS {db_id} - memory appears constrained")
                    continue
            
            # Skip if high IOPS activity (>1000 peak) - may be I/O bound
            if max_read_iops > 1000 or max_write_iops > 1000:
                print(f"Skipping RDS {db_id} - high I/O activity")
                continue
            
            # CONSERVATIVE thresholds for RDS - databases are critical infrastructure
            # Only recommend downsizing if utilization is extremely low over 14 days
            # AND peak connections are very low (indicating truly unused capacity)
            if avg_cpu < 10 and max_cpu < 25 and avg_conn < 3 and max_conn < 10:
                try:
                    current_cost = get_rds_cost(db_class, engine, session.region_name, multi_az)
                    smaller_class = get_smaller_rds_class(db_class)
                    
                    if smaller_class:
                        smaller_cost = get_rds_cost(smaller_class, engine, session.region_name, multi_az)
                        monthly_savings = (current_cost - smaller_cost) * 730
                        
                        if monthly_savings > 0:
                            # Get tags
                            rds_tags = {}
                            try:
                                tag_response = rds.list_tags_for_resource(ResourceName=db['DBInstanceArn'])
                                rds_tags = get_resource_tags(tag_response.get('TagList', []))
                            except Exception:
                                pass
                            
                            recommendations.append({
                                'db_id': db_id,
                                'current_class': db_class,
                                'recommended_class': smaller_class,
                                'engine': engine,
                                'current_cost': round(current_cost * 730, 2),
                                'recommended_cost': round(smaller_cost * 730, 2),
                                'monthly_savings': round(monthly_savings, 2),
                                'reason': f'Very low utilization (CPU avg: {avg_cpu:.1f}%, max: {max_cpu:.1f}%, Connections avg: {avg_conn:.0f}, max: {max_conn:.0f})',
                                'confidence': 'Medium',
                                'data_points': len(cpu_stats['Datapoints']),
                                'tags': rds_tags
                            })
                except PricingUnavailableError as e:
                    skipped_resources.append(f"RDS {db_id}: {e}")
                    print(f"Skipping RDS {db_id} - pricing unavailable: {e}")
    except Exception as e:
        print(f"RDS scan error: {e}")
    
    return recommendations


def scan_lambda_functions(session):
    recommendations = []
    skipped_resources = []
    lambda_client = session.client('lambda')
    cloudwatch = session.client('cloudwatch')
    
    # Minimum data points required for reliable analysis (at least 7 days of data)
    MIN_DATA_POINTS = 7
    
    try:
        # Get all Lambda functions with pagination
        all_functions = []
        paginator = lambda_client.get_paginator('list_functions')
        for page in paginator.paginate():
            all_functions.extend(page['Functions'])
        
        for func in all_functions:
            func_name = func['FunctionName']
            memory_size = func['MemorySize']
            
            # Get metrics over 14 days
            end_time = datetime.now(timezone.utc)
            start_time = end_time - timedelta(days=14)
            
            duration_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/Lambda',
                MetricName='Duration',
                Dimensions=[{'Name': 'FunctionName', 'Value': func_name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            invocations = cloudwatch.get_metric_statistics(
                Namespace='AWS/Lambda',
                MetricName='Invocations',
                Dimensions=[{'Name': 'FunctionName', 'Value': func_name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            # Check for errors - don't recommend changes to functions with high error rates
            errors = cloudwatch.get_metric_statistics(
                Namespace='AWS/Lambda',
                MetricName='Errors',
                Dimensions=[{'Name': 'FunctionName', 'Value': func_name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            # Check for throttles - may indicate function is already under pressure
            throttles = cloudwatch.get_metric_statistics(
                Namespace='AWS/Lambda',
                MetricName='Throttles',
                Dimensions=[{'Name': 'FunctionName', 'Value': func_name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            # Require minimum data points for reliable analysis
            if not duration_stats['Datapoints'] or len(duration_stats['Datapoints']) < MIN_DATA_POINTS:
                continue
            
            if not invocations['Datapoints'] or len(invocations['Datapoints']) < MIN_DATA_POINTS:
                continue
            
            avg_duration = sum(d['Average'] for d in duration_stats['Datapoints']) / len(duration_stats['Datapoints'])
            max_duration = max(d['Maximum'] for d in duration_stats['Datapoints'])
            total_invocations = sum(d['Sum'] for d in invocations['Datapoints'])
            
            # Calculate error and throttle totals
            total_errors = sum(d['Sum'] for d in errors['Datapoints']) if errors['Datapoints'] else 0
            total_throttles = sum(d['Sum'] for d in throttles['Datapoints']) if throttles['Datapoints'] else 0
            
            # Skip if no meaningful invocations
            if total_invocations < 100:
                continue
            
            # Calculate error rate - skip if > 1% error rate
            error_rate = (total_errors / total_invocations) * 100 if total_invocations > 0 else 0
            if error_rate > 1:
                print(f"Skipping Lambda {func_name} - high error rate ({error_rate:.2f}%)")
                continue
            
            # Skip if there were any throttles - indicates potential capacity issues
            if total_throttles > 0:
                print(f"Skipping Lambda {func_name} - throttles detected ({total_throttles})")
                continue
            
            # CONSERVATIVE Lambda memory recommendations
            # Only recommend reduction if:
            # 1. Memory is significantly over-provisioned (> 1024 MB)
            # 2. Average duration is very short (< 500ms)
            # 3. Max duration is also low (< 2000ms) - ensures headroom for cold starts and spikes
            # 4. Only reduce by 25% (not 50%) to maintain buffer
            # 5. No errors or throttles
            if memory_size > 1024 and avg_duration < 500 and max_duration < 2000:
                # Conservative: only reduce by 25%, not 50%
                recommended_memory = max(256, int(memory_size * 0.75))
                
                try:
                    current_cost = calculate_lambda_cost(memory_size, avg_duration, total_invocations, session.region_name)
                    recommended_cost = calculate_lambda_cost(recommended_memory, avg_duration, total_invocations, session.region_name)
                    monthly_savings = current_cost - recommended_cost
                    
                    if monthly_savings > 5:  # Only recommend if savings > $5/month (more meaningful threshold)
                        # Get function tags
                        func_tags = {}
                        try:
                            tag_response = lambda_client.list_tags(Resource=func['FunctionArn'])
                            func_tags = get_resource_tags([{'Key': k, 'Value': v} for k, v in tag_response.get('Tags', {}).items()])
                        except Exception:
                            pass
                        
                        recommendations.append({
                            'function_name': func_name,
                            'current_memory': memory_size,
                            'recommended_memory': recommended_memory,
                            'avg_duration': round(avg_duration, 0),
                            'max_duration': round(max_duration, 0),
                            'invocations': int(total_invocations),
                            'error_rate': round(error_rate, 2),
                            'current_cost': round(current_cost, 2),
                            'recommended_cost': round(recommended_cost, 2),
                            'monthly_savings': round(monthly_savings, 2),
                            'confidence': 'Medium',
                            'data_points': len(duration_stats['Datapoints']),
                            'tags': func_tags
                        })
                except PricingUnavailableError as e:
                    skipped_resources.append(f"Lambda {func_name}: {e}")
                    print(f"Skipping Lambda {func_name} - pricing unavailable: {e}")
    except Exception as e:
        print(f"Lambda scan error: {e}")
    
    return recommendations


def scan_elastic_ips(session):
    recommendations = []
    skipped_resources = []
    ec2 = session.client('ec2')
    
    try:
        addresses = ec2.describe_addresses()
        
        try:
            # Real-time public-IPv4 hourly price. Since 2024-02-01 AWS bills every
            # public IPv4 (~$0.005/hr = ~$3.65/mo) whether attached or not, so an
            # unattached EIP is pure waste.
            eip_hourly_cost = get_eip_cost(session.region_name)
            
            for addr in addresses['Addresses']:
                # Check if unattached (handle both VPC and EC2-Classic scenarios)
                is_attached = 'InstanceId' in addr or 'NetworkInterfaceId' in addr
                if not is_attached:
                    # Use AllocationId if available (VPC), otherwise use PublicIp as identifier
                    allocation_id = addr.get('AllocationId', addr.get('PublicIp', 'N/A'))
                    monthly_cost = eip_hourly_cost * 730  # hours per month
                    
                    recommendations.append({
                        'ip_address': addr['PublicIp'],
                        'allocation_id': allocation_id,
                        'status': 'Unattached',
                        'monthly_savings': round(monthly_cost, 2),
                        'recommendation': 'Release this idle public IPv4 address (billed even while unattached since Feb 2024)',
                        'confidence': 'High',
                        'tags': get_resource_tags(addr.get('Tags', []))
                    })

            # Public IPv4 footprint: count ALL public IPv4 in use (EIP + auto-assigned)
            # so the report can surface the post-2024 charge as an optimization lever.
            try:
                in_use_public_ips = 0
                ni_paginator = ec2.get_paginator('describe_network_interfaces')
                for page in ni_paginator.paginate():
                    for ni in page.get('NetworkInterfaces', []):
                        if ni.get('Association', {}).get('PublicIp'):
                            in_use_public_ips += 1
                if in_use_public_ips > 0:
                    monthly_ipv4_cost = in_use_public_ips * eip_hourly_cost * 730
                    recommendations.append({
                        'ip_address': f'{in_use_public_ips} in-use public IPv4 address(es)',
                        'allocation_id': 'N/A',
                        'status': 'In use',
                        'count': in_use_public_ips,
                        'monthly_cost': round(monthly_ipv4_cost, 2),
                        # Informational: reducing in-use IPs needs architecture change,
                        # so we do not claim it as directly reclaimable savings.
                        'monthly_savings': 0.0,
                        'recommendation': 'Reduce public IPv4 usage (private subnets + VPC endpoints/NAT, IPv6, or BYOIP) to cut the $0.005/hr per-IP charge',
                        'confidence': 'Low',
                        'effort': 'Medium',
                        'risk': 'Low',
                        'tags': {}
                    })
            except Exception as e:
                print(f"Public IPv4 inventory warning: {e}")
        except PricingUnavailableError as e:
            skipped_resources.append(f"EIP pricing: {e}")
            print(f"Skipping all EIPs - pricing unavailable: {e}")
    except Exception as e:
        print(f"EIP scan error: {e}")
    
    return recommendations


def scan_s3_buckets(session):
    """Scan S3 buckets for optimization opportunities."""
    recommendations = []
    s3 = session.client('s3')
    
    try:
        buckets = s3.list_buckets().get('Buckets', [])
        
        for bucket in buckets:
            bucket_name = bucket['Name']
            
            # Get bucket region
            try:
                location = s3.get_bucket_location(Bucket=bucket_name)
                bucket_region = location.get('LocationConstraint') or 'us-east-1'
            except Exception:
                bucket_region = 'unknown'
            
            # Get tags
            tags = {}
            try:
                tag_response = s3.get_bucket_tagging(Bucket=bucket_name)
                tags = get_resource_tags(tag_response.get('TagSet', []))
            except Exception:
                pass
            
            issues = []
            
            # Check lifecycle policy
            has_lifecycle = False
            try:
                s3.get_bucket_lifecycle_configuration(Bucket=bucket_name)
                has_lifecycle = True
            except Exception:
                pass
            if not has_lifecycle:
                issues.append('No lifecycle policy configured')
            
            # Check Intelligent-Tiering
            has_intelligent_tiering = False
            try:
                configs = s3.list_bucket_intelligent_tiering_configurations(Bucket=bucket_name)
                if configs.get('IntelligentTieringConfigurationList'):
                    has_intelligent_tiering = True
            except Exception:
                pass
            if not has_intelligent_tiering:
                issues.append('No Intelligent-Tiering configured')
            
            # Check for incomplete multipart uploads
            incomplete_uploads = 0
            try:
                response = s3.list_multipart_uploads(Bucket=bucket_name)
                incomplete_uploads = len(response.get('Uploads', []))
            except Exception:
                pass
            if incomplete_uploads > 0:
                issues.append(f'{incomplete_uploads} incomplete multipart upload(s)')
            
            if issues:
                recommendations.append({
                    'bucket_name': bucket_name,
                    'region': bucket_region,
                    'issues': ', '.join(issues),
                    'has_lifecycle': has_lifecycle,
                    'has_intelligent_tiering': has_intelligent_tiering,
                    'incomplete_uploads': incomplete_uploads,
                    'recommendation': '; '.join([
                        'Add lifecycle policy to transition/expire objects' if not has_lifecycle else '',
                        'Enable Intelligent-Tiering for automatic cost optimization' if not has_intelligent_tiering else '',
                        f'Abort {incomplete_uploads} incomplete multipart upload(s) to reclaim storage' if incomplete_uploads > 0 else ''
                    ]).strip('; '),
                    'monthly_savings': 0.0,
                    'confidence': 'Medium',
                    'tags': tags
                })
    except Exception as e:
        print(f"S3 scan error: {e}")
    
    return recommendations


def scan_stopped_ec2_instances(session):
    """Scan for long-stopped EC2 instances with attached EBS volumes."""
    recommendations = []
    skipped_resources = []
    ec2 = session.client('ec2')
    
    try:
        # Get all stopped instances with pagination
        all_instances = []
        paginator = ec2.get_paginator('describe_instances')
        for page in paginator.paginate(Filters=[{'Name': 'instance-state-name', 'Values': ['stopped']}]):
            for reservation in page['Reservations']:
                all_instances.extend(reservation['Instances'])
        
        for instance in all_instances:
            instance_id = instance['InstanceId']
            instance_type = instance['InstanceType']
            
            # Parse stop time from StateTransitionReason
            # Format: "User initiated (2024-01-15 10:30:00 GMT)"
            stopped_days = None
            reason = instance.get('StateTransitionReason', '')
            try:
                if '(' in reason and ')' in reason:
                    date_str = reason.split('(')[1].split(')')[0].replace(' GMT', '')
                    stop_time = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                    stop_time = stop_time.replace(tzinfo=timezone.utc)
                    stopped_days = (datetime.now(timezone.utc) - stop_time).days
            except Exception:
                pass
            
            # Only flag instances stopped for 30+ days
            if stopped_days is not None and stopped_days < 30:
                continue
            
            # Calculate EBS cost for attached volumes
            total_ebs_cost = 0.0
            attached_volumes = []
            
            for mapping in instance.get('BlockDeviceMappings', []):
                volume_id = mapping.get('Ebs', {}).get('VolumeId')
                if volume_id:
                    try:
                        vol_response = ec2.describe_volumes(VolumeIds=[volume_id])
                        if vol_response['Volumes']:
                            vol = vol_response['Volumes'][0]
                            vol_cost = calculate_ebs_cost(
                                vol['VolumeType'], vol['Size'], session.region_name,
                                vol.get('Iops', 0), vol.get('Throughput', 0)
                            )
                            total_ebs_cost += vol_cost
                            attached_volumes.append({
                                'volume_id': volume_id,
                                'size': vol['Size'],
                                'type': vol['VolumeType'],
                                'monthly_cost': round(vol_cost, 2)
                            })
                    except Exception as e:
                        print(f"Error checking volume {volume_id}: {e}")
            
            if total_ebs_cost > 0:
                days_str = f'{stopped_days} days' if stopped_days is not None else '30+ days'
                recommendations.append({
                    'instance_id': instance_id,
                    'instance_type': instance_type,
                    'stopped_days': stopped_days or 30,
                    'attached_volumes': len(attached_volumes),
                    'monthly_savings': round(total_ebs_cost, 2),
                    'reason': f'Instance stopped for {days_str} with {len(attached_volumes)} attached EBS volume(s)',
                    'recommendation': 'Create AMI backup and terminate instance, or delete unneeded EBS volumes',
                    'confidence': 'High',
                    'tags': get_resource_tags(instance.get('Tags', []))
                })
    except Exception as e:
        print(f"Stopped EC2 scan error: {e}")
    
    return recommendations


def scan_nat_gateways(session):
    """Scan NAT Gateways for underutilization."""
    recommendations = []
    ec2 = session.client('ec2')
    cloudwatch = session.client('cloudwatch')

    # Live NAT gateway pricing (falls back to common us-east-1 rates).
    nat_pricing = get_nat_gateway_pricing(session.region_name)
    NAT_GW_HOURLY_COST = nat_pricing['hourly']
    NAT_GW_DATA_COST_PER_GB = nat_pricing['per_gb']
    
    try:
        nat_gateways = ec2.describe_nat_gateways(
            Filter=[{'Name': 'state', 'Values': ['available']}]
        ).get('NatGateways', [])
        
        for nat_gw in nat_gateways:
            nat_gw_id = nat_gw['NatGatewayId']
            
            # Check data processed in last 14 days
            end_time = datetime.now(timezone.utc)
            start_time = end_time - timedelta(days=14)
            
            bytes_out = cloudwatch.get_metric_statistics(
                Namespace='AWS/NATGateway',
                MetricName='BytesOutToDestination',
                Dimensions=[{'Name': 'NatGatewayId', 'Value': nat_gw_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            bytes_in = cloudwatch.get_metric_statistics(
                Namespace='AWS/NATGateway',
                MetricName='BytesInFromSource',
                Dimensions=[{'Name': 'NatGatewayId', 'Value': nat_gw_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            active_conn = cloudwatch.get_metric_statistics(
                Namespace='AWS/NATGateway',
                MetricName='ActiveConnectionCount',
                Dimensions=[{'Name': 'NatGatewayId', 'Value': nat_gw_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Average', 'Maximum']
            )
            
            total_bytes_out = sum(d['Sum'] for d in bytes_out['Datapoints']) if bytes_out['Datapoints'] else 0
            total_bytes_in = sum(d['Sum'] for d in bytes_in['Datapoints']) if bytes_in['Datapoints'] else 0
            total_gb = (total_bytes_in + total_bytes_out) / (1024**3)
            avg_daily_gb = total_gb / 14 if total_gb > 0 else 0
            
            avg_connections = 0
            if active_conn['Datapoints']:
                avg_connections = sum(d['Average'] for d in active_conn['Datapoints']) / len(active_conn['Datapoints'])
            
            monthly_base_cost = NAT_GW_HOURLY_COST * 730
            monthly_data_cost = avg_daily_gb * 30 * NAT_GW_DATA_COST_PER_GB
            total_monthly_cost = monthly_base_cost + monthly_data_cost
            
            # Flag if very low data processed (< 1 GB/day avg)
            if avg_daily_gb < 1:
                tags = get_resource_tags(nat_gw.get('Tags', []))
                
                recommendations.append({
                    'nat_gateway_id': nat_gw_id,
                    'vpc_id': nat_gw.get('VpcId', 'N/A'),
                    'subnet_id': nat_gw.get('SubnetId', 'N/A'),
                    'state': nat_gw.get('State', 'N/A'),
                    'avg_daily_gb': round(avg_daily_gb, 2),
                    'avg_connections': round(avg_connections, 1),
                    'monthly_cost': round(total_monthly_cost, 2),
                    'monthly_savings': round(total_monthly_cost, 2),
                    'reason': f'Low data transfer ({avg_daily_gb:.2f} GB/day avg, {avg_connections:.0f} avg connections)',
                    'recommendation': 'Consider removing if not needed, or use VPC endpoints for AWS service traffic',
                    'confidence': 'Medium',
                    'tags': tags
                })
    except Exception as e:
        print(f"NAT Gateway scan error: {e}")
    
    return recommendations


def scan_dynamodb_tables(session):
    """Scan DynamoDB tables for optimization opportunities."""
    recommendations = []
    dynamodb = session.client('dynamodb')
    cloudwatch = session.client('cloudwatch')
    
    MIN_DATA_POINTS = 7
    
    try:
        # Get all tables with pagination
        tables = []
        paginator = dynamodb.get_paginator('list_tables')
        for page in paginator.paginate():
            tables.extend(page['TableNames'])
        
        for table_name in tables:
            table = dynamodb.describe_table(TableName=table_name)['Table']
            billing_mode = table.get('BillingModeSummary', {}).get('BillingMode', 'PROVISIONED')
            
            if billing_mode != 'PROVISIONED':
                continue
            
            provisioned_rcu = table['ProvisionedThroughput']['ReadCapacityUnits']
            provisioned_wcu = table['ProvisionedThroughput']['WriteCapacityUnits']
            
            if provisioned_rcu == 0 and provisioned_wcu == 0:
                continue
            
            # Check CloudWatch for actual usage over 14 days
            end_time = datetime.now(timezone.utc)
            start_time = end_time - timedelta(days=14)
            
            consumed_rcu_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/DynamoDB',
                MetricName='ConsumedReadCapacityUnits',
                Dimensions=[{'Name': 'TableName', 'Value': table_name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            consumed_wcu_stats = cloudwatch.get_metric_statistics(
                Namespace='AWS/DynamoDB',
                MetricName='ConsumedWriteCapacityUnits',
                Dimensions=[{'Name': 'TableName', 'Value': table_name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=86400,
                Statistics=['Sum']
            )
            
            if not consumed_rcu_stats['Datapoints'] or len(consumed_rcu_stats['Datapoints']) < MIN_DATA_POINTS:
                continue
            if not consumed_wcu_stats['Datapoints'] or len(consumed_wcu_stats['Datapoints']) < MIN_DATA_POINTS:
                continue
            
            # ConsumedReadCapacityUnits is a SUM metric (total units consumed per
            # period). The previous code took Statistic='Average' and multiplied by
            # 86400 seconds, which is dimensionally wrong. Use Sum for the true
            # total consumed over the window, then derive a per-second rate for
            # the utilization comparison.
            window_seconds = (end_time - start_time).total_seconds()
            total_rcu = sum(d['Sum'] for d in consumed_rcu_stats['Datapoints'])
            total_wcu = sum(d['Sum'] for d in consumed_wcu_stats['Datapoints'])
            avg_rcu = total_rcu / window_seconds if window_seconds else 0  # consumed RCU/sec
            avg_wcu = total_wcu / window_seconds if window_seconds else 0  # consumed WCU/sec
            
            rcu_utilization = (avg_rcu / provisioned_rcu * 100) if provisioned_rcu > 0 else 0
            wcu_utilization = (avg_wcu / provisioned_wcu * 100) if provisioned_wcu > 0 else 0
            
            # Flag if both RCU and WCU utilization are very low
            if rcu_utilization < 20 and wcu_utilization < 20:
                # Live DynamoDB pricing (falls back to current us-east-1 rates).
                ddb_pricing = get_dynamodb_pricing(session.region_name)
                provisioned_monthly = (provisioned_rcu * ddb_pricing['rcu_hour']
                                       + provisioned_wcu * ddb_pricing['wcu_hour']) * 730
                
                # On-demand cost = total request units scaled to a month x per-request price.
                monthly_rru = total_rcu / window_seconds * 730 * 3600 if window_seconds else 0
                monthly_wru = total_wcu / window_seconds * 730 * 3600 if window_seconds else 0
                on_demand_monthly = (monthly_rru * ddb_pricing['rru_per_million'] / 1_000_000
                                     + monthly_wru * ddb_pricing['wru_per_million'] / 1_000_000)
                
                monthly_savings = provisioned_monthly - on_demand_monthly
                
                if monthly_savings > 1:
                    # Get tags
                    table_tags = {}
                    try:
                        tag_response = dynamodb.list_tags_of_resource(ResourceArn=table['TableArn'])
                        table_tags = get_resource_tags(tag_response.get('Tags', []))
                    except Exception:
                        pass
                    
                    recommendations.append({
                        'table_name': table_name,
                        'billing_mode': billing_mode,
                        'provisioned_rcu': provisioned_rcu,
                        'provisioned_wcu': provisioned_wcu,
                        'avg_rcu': round(avg_rcu, 2),
                        'avg_wcu': round(avg_wcu, 2),
                        'rcu_utilization': round(rcu_utilization, 1),
                        'wcu_utilization': round(wcu_utilization, 1),
                        'current_cost': round(provisioned_monthly, 2),
                        'recommended_cost': round(on_demand_monthly, 2),
                        'monthly_savings': round(monthly_savings, 2),
                        'pricing_source': ddb_pricing['source'],
                        'recommendation': 'Switch to On-Demand billing mode (or lower provisioned capacity / add auto scaling)',
                        'reason': f'Low utilization (RCU: {rcu_utilization:.1f}%, WCU: {wcu_utilization:.1f}%)',
                        'confidence': 'Medium',
                        'tags': table_tags
                    })
    except Exception as e:
        print(f"DynamoDB scan error: {e}")
    
    return recommendations


def get_ebs_snapshot_price(region):
    """Get EBS snapshot storage price ($/GB-month) from Price List (cached, fallback)."""
    cache_key = f"ebs_snapshot_{region}"
    if cache_key in PRICING_CACHE:
        cached = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached['timestamp'] < CACHE_TTL:
            return cached['price']
    price = 0.05  # us-east-1 standard snapshot fallback
    try:
        location = REGION_LOCATION_MAP.get(region)
        if location:
            pricing_client = boto3.client('pricing', region_name='us-east-1')
            response = pricing_client.get_products(
                ServiceCode='AmazonEC2',
                Filters=[
                    {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'Storage Snapshot'},
                    {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                ],
                MaxResults=20
            )
            for price_item in response.get('PriceList', []):
                data = json.loads(price_item)
                on_demand = data.get('terms', {}).get('OnDemand', {})
                if not on_demand:
                    continue
                for dim in list(list(on_demand.values())[0]['priceDimensions'].values()):
                    if 'GB-Mo' in dim.get('unit', ''):
                        usd = dim.get('pricePerUnit', {}).get('USD')
                        if usd and float(usd) > 0:
                            price = float(usd)
                            break
    except Exception as e:
        print(f"Snapshot pricing fell back to default: {e}")
    PRICING_CACHE[cache_key] = {'price': price, 'timestamp': datetime.now().timestamp()}
    return price


def get_load_balancer_price(region, lb_type):
    """Get load balancer hourly base price from Price List (cached, fallback).

    Returns the fixed hourly charge (excludes LCU/hour, which is usage based).
    """
    lb_type = (lb_type or 'application').lower()
    cache_key = f"elb_{lb_type}_{region}"
    if cache_key in PRICING_CACHE:
        cached = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached['timestamp'] < CACHE_TTL:
            return cached['price']
    # us-east-1 fallbacks
    fallback = {'application': 0.0225, 'network': 0.0225, 'gateway': 0.0125, 'classic': 0.025}
    price = fallback.get(lb_type, 0.0225)
    product_family = {
        'application': 'Load Balancer-Application',
        'network': 'Load Balancer-Network',
        'gateway': 'Load Balancer-Gateway',
        'classic': 'Load Balancer',
    }.get(lb_type, 'Load Balancer-Application')
    try:
        location = REGION_LOCATION_MAP.get(region)
        if location:
            pricing_client = boto3.client('pricing', region_name='us-east-1')
            response = pricing_client.get_products(
                ServiceCode='AWSELB',
                Filters=[
                    {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': product_family},
                    {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                ],
                MaxResults=20
            )
            for price_item in response.get('PriceList', []):
                data = json.loads(price_item)
                on_demand = data.get('terms', {}).get('OnDemand', {})
                if not on_demand:
                    continue
                for dim in list(list(on_demand.values())[0]['priceDimensions'].values()):
                    unit = dim.get('unit', '').lower()
                    usd = dim.get('pricePerUnit', {}).get('USD')
                    if usd and float(usd) > 0 and ('hrs' in unit or 'hour' in unit):
                        price = float(usd)
                        break
    except Exception as e:
        print(f"Load balancer pricing fell back to default: {e}")
    PRICING_CACHE[cache_key] = {'price': price, 'timestamp': datetime.now().timestamp()}
    return price


def scan_ebs_snapshots(session):
    """Flag orphaned or old EBS snapshots. Each snapshot is billed per GB-month.

    Snapshot storage is incremental, so per-snapshot size is an UPPER BOUND on
    reclaimable cost; we mark it as estimated and only flag snapshots that are
    orphaned (source volume gone) or old, and never those backing an AMI.
    """
    recommendations = []
    ec2 = session.client('ec2')
    snap_price = get_ebs_snapshot_price(session.region_name)

    try:
        # Snapshots referenced by self-owned AMIs must not be deleted.
        ami_snapshot_ids = set()
        try:
            for img in ec2.describe_images(Owners=['self']).get('Images', []):
                for bdm in img.get('BlockDeviceMappings', []):
                    sid = bdm.get('Ebs', {}).get('SnapshotId')
                    if sid:
                        ami_snapshot_ids.add(sid)
        except Exception as e:
            print(f"AMI cross-reference warning: {e}")

        # Existing volumes (to detect orphaned snapshots).
        existing_volumes = set()
        try:
            vp = ec2.get_paginator('describe_volumes')
            for page in vp.paginate():
                for v in page['Volumes']:
                    existing_volumes.add(v['VolumeId'])
        except Exception as e:
            print(f"Volume list warning: {e}")

        now = datetime.now(timezone.utc)
        paginator = ec2.get_paginator('describe_snapshots')
        for page in paginator.paginate(OwnerIds=['self']):
            for snap in page.get('Snapshots', []):
                snap_id = snap['SnapshotId']
                if snap_id in ami_snapshot_ids:
                    continue
                size = snap.get('VolumeSize', 0)
                start = snap.get('StartTime')
                age_days = (now - start).days if start else 0
                source_vol = snap.get('VolumeId', '')
                orphaned = bool(source_vol) and source_vol not in existing_volumes

                if not orphaned and age_days <= 90:
                    continue

                monthly_cost = size * snap_price
                if monthly_cost <= 0:
                    continue

                recommendations.append({
                    'snapshot_id': snap_id,
                    'size': size,
                    'age_days': age_days,
                    'source_volume': source_vol or 'N/A',
                    'issue': 'Orphaned' if orphaned else 'Old',
                    'monthly_savings': round(monthly_cost, 2),
                    'recommendation': ('Delete orphaned snapshot (source volume no longer exists)'
                                       if orphaned else
                                       f'Review/delete snapshot older than 90 days ({age_days}d)') +
                                      ' - estimate is an upper bound (snapshots are incremental)',
                    'confidence': 'High' if orphaned else 'Medium',
                    'effort': 'Low',
                    'risk': 'Medium',
                    'tags': get_resource_tags(snap.get('Tags', []))
                })
    except Exception as e:
        print(f"EBS snapshot scan error: {e}")

    return recommendations


def scan_load_balancers(session):
    """Flag idle load balancers (ALB/NLB/GWLB via ELBv2, plus Classic ELB).

    A load balancer with ~zero traffic over 14 days is paying its hourly base
    charge for nothing.
    """
    recommendations = []
    cloudwatch = session.client('cloudwatch')
    end_time = datetime.now(timezone.utc)
    start_time = end_time - timedelta(days=14)

    # ALB / NLB / GWLB
    try:
        elbv2 = session.client('elbv2')
        paginator = elbv2.get_paginator('describe_load_balancers')
        for page in paginator.paginate():
            for lb in page.get('LoadBalancers', []):
                lb_arn = lb['LoadBalancerArn']
                lb_name = lb['LoadBalancerName']
                lb_type = lb.get('Type', 'application')
                # Don't flag brand-new load balancers (may not have traffic yet).
                created = lb.get('CreatedTime')
                if created and (datetime.now(timezone.utc) - created).days < 7:
                    continue
                # CloudWatch dimension value is the ARN suffix after 'loadbalancer/'.
                try:
                    dim_value = lb_arn.split('loadbalancer/')[1]
                except IndexError:
                    continue

                if lb_type == 'application':
                    namespace, metric, stat = 'AWS/ApplicationELB', 'RequestCount', 'Sum'
                elif lb_type == 'network':
                    namespace, metric, stat = 'AWS/NetworkELB', 'ProcessedBytes', 'Sum'
                else:
                    namespace, metric, stat = 'AWS/GatewayELB', 'ProcessedBytes', 'Sum'

                stats = cloudwatch.get_metric_statistics(
                    Namespace=namespace, MetricName=metric,
                    Dimensions=[{'Name': 'LoadBalancer', 'Value': dim_value}],
                    StartTime=start_time, EndTime=end_time, Period=86400, Statistics=[stat]
                )
                total = sum(d[stat] for d in stats['Datapoints']) if stats['Datapoints'] else 0
                # Idle if effectively no traffic across 14 days.
                if total < 1:
                    price = get_load_balancer_price(session.region_name, lb_type)
                    monthly = price * 730
                    tags = {}
                    try:
                        tag_desc = elbv2.describe_tags(ResourceArns=[lb_arn])
                        tags = get_resource_tags(tag_desc['TagDescriptions'][0].get('Tags', []))
                    except Exception:
                        pass
                    recommendations.append({
                        'load_balancer_name': lb_name,
                        'load_balancer_arn': lb_arn,
                        'type': lb_type,
                        'metric': f'{metric}={int(total)} over 14d',
                        'monthly_savings': round(monthly, 2),
                        'reason': f'No traffic in 14 days ({metric} ~ 0)',
                        'recommendation': 'Delete idle load balancer if no longer needed',
                        'confidence': 'High',
                        'effort': 'Low',
                        'risk': 'Medium',
                        'tags': tags
                    })
    except Exception as e:
        print(f"ELBv2 scan error: {e}")

    # Classic ELB
    try:
        elb = session.client('elb')
        paginator = elb.get_paginator('describe_load_balancers')
        for page in paginator.paginate():
            for lb in page.get('LoadBalancerDescriptions', []):
                lb_name = lb['LoadBalancerName']
                stats = cloudwatch.get_metric_statistics(
                    Namespace='AWS/ELB', MetricName='RequestCount',
                    Dimensions=[{'Name': 'LoadBalancerName', 'Value': lb_name}],
                    StartTime=start_time, EndTime=end_time, Period=86400, Statistics=['Sum']
                )
                total = sum(d['Sum'] for d in stats['Datapoints']) if stats['Datapoints'] else 0
                if total < 1:
                    price = get_load_balancer_price(session.region_name, 'classic')
                    monthly = price * 730
                    recommendations.append({
                        'load_balancer_name': lb_name,
                        'type': 'classic',
                        'metric': f'RequestCount={int(total)} over 14d',
                        'monthly_savings': round(monthly, 2),
                        'reason': 'No requests in 14 days',
                        'recommendation': 'Delete idle Classic Load Balancer if no longer needed',
                        'confidence': 'High',
                        'effort': 'Low',
                        'risk': 'Medium',
                        'tags': {}
                    })
    except Exception as e:
        print(f"Classic ELB scan error: {e}")

    return recommendations


def scan_savings_plans_purchase(session, term='ONE_YEAR', payment='NO_UPFRONT',
                                lookback='THIRTY_DAYS'):
    """Get Compute Savings Plans purchase recommendation from Cost Explorer.

    This is RATE optimization (commitment discount), distinct from waste
    elimination. Savings are labeled savings_basis='commitment_purchase' so the
    report never conflates them with rightsizing/idle savings (double counting).
    """
    recommendations = []
    try:
        ce = session.client('ce', region_name='us-east-1')
        resp = ce.get_savings_plans_purchase_recommendation(
            SavingsPlansType='COMPUTE_SP',
            TermInYears=term,
            PaymentOption=payment,
            LookbackPeriodInDays=lookback,
        )
        detail = resp.get('SavingsPlansPurchaseRecommendation', {})
        summary = detail.get('SavingsPlansPurchaseRecommendationSummary', {})
        est = float(summary.get('EstimatedMonthlySavingsAmount', 0) or 0)
        if est > 0:
            recommendations.append({
                'type': 'Compute Savings Plan',
                'term': term.replace('_', ' ').title(),
                'payment_option': payment.replace('_', ' ').title(),
                'hourly_commitment': summary.get('HourlyCommitmentToPurchase', 'N/A'),
                'estimated_savings_pct': summary.get('EstimatedSavingsPercentage', 'N/A'),
                'current_cost': round(float(summary.get('CurrentOnDemandSpend', 0) or 0), 2),
                'monthly_savings': round(est, 2),
                'savings_basis': 'commitment_purchase',
                'reason': f"~{summary.get('EstimatedSavingsPercentage', '?')}% off analyzed on-demand compute spend",
                'recommendation': (f"Purchase a {term.replace('_', ' ').lower()} "
                                   f"{payment.replace('_', ' ').lower()} Compute Savings Plan"),
                'confidence': 'High',
                'effort': 'Low',
                'risk': 'Low',
                'tags': {}
            })
    except Exception as e:
        print(f"Savings Plans purchase recommendation unavailable: {e}")
    return recommendations


def get_cost_forecast_and_spend(session):
    """Return {'month_to_date': x, 'forecast_month': y} from Cost Explorer (best effort)."""
    result = {'month_to_date': None, 'forecast_month': None}
    try:
        ce = session.client('ce', region_name='us-east-1')
        today = datetime.now(timezone.utc).date()
        first_of_month = today.replace(day=1)

        # Month-to-date actual spend.
        try:
            if today > first_of_month:
                mtd = ce.get_cost_and_usage(
                    TimePeriod={'Start': first_of_month.strftime('%Y-%m-%d'),
                                'End': today.strftime('%Y-%m-%d')},
                    Granularity='MONTHLY', Metrics=['UnblendedCost']
                )
                results = mtd.get('ResultsByTime', [])
                if results:
                    amt = results[0].get('Total', {}).get('UnblendedCost', {}).get('Amount')
                    if amt is not None:
                        result['month_to_date'] = round(float(amt), 2)
        except Exception as e:
            print(f"MTD spend unavailable: {e}")

        # Forecast to end of month.
        try:
            end = (first_of_month + timedelta(days=32)).replace(day=1)
            if end > today:
                fc = ce.get_cost_forecast(
                    TimePeriod={'Start': today.strftime('%Y-%m-%d'), 'End': end.strftime('%Y-%m-%d')},
                    Metric='UNBLENDED_COST', Granularity='MONTHLY'
                )
                amt = fc.get('Total', {}).get('Amount')
                if amt is not None:
                    result['forecast_month'] = round(float(amt), 2)
        except Exception as e:
            print(f"Cost forecast unavailable: {e}")
    except Exception as e:
        print(f"Cost Explorer forecast/spend skipped: {e}")
    return result


def scan_ri_sp_coverage(session):
    """Scan Reserved Instance and Savings Plans coverage."""
    summary = {
        'total_running_instances': 0,
        'ri_covered_instances': 0,
        'ri_coverage_pct': 0.0,
        'active_ris': [],
        'savings_plans': [],
        'sp_coverage_pct': 0.0
    }
    
    ec2 = session.client('ec2')
    
    # Count running instances
    try:
        paginator = ec2.get_paginator('describe_instances')
        for page in paginator.paginate(Filters=[{'Name': 'instance-state-name', 'Values': ['running']}]):
            for reservation in page['Reservations']:
                summary['total_running_instances'] += len(reservation['Instances'])
    except Exception as e:
        print(f"Error counting instances: {e}")
    
    # Get active Reserved Instances
    try:
        ris = ec2.describe_reserved_instances(Filters=[{'Name': 'state', 'Values': ['active']}])
        total_ri_count = 0
        for ri in ris['ReservedInstances']:
            count = ri['InstanceCount']
            total_ri_count += count
            summary['active_ris'].append({
                'instance_type': ri['InstanceType'],
                'count': count,
                'offering_type': ri.get('OfferingType', 'N/A'),
                'end_date': ri.get('End').strftime('%Y-%m-%d') if isinstance(ri.get('End'), datetime) else str(ri.get('End', 'N/A'))
            })
        summary['ri_covered_instances'] = total_ri_count
        if summary['total_running_instances'] > 0:
            summary['ri_coverage_pct'] = round(total_ri_count / summary['total_running_instances'] * 100, 1)
    except Exception as e:
        print(f"RI check error: {e}")
    
    # Get Savings Plans
    try:
        sp_client = session.client('savingsplans')
        sp_response = sp_client.describe_savings_plans(
            states=['active']
        )
        for sp in sp_response.get('savingsPlans', []):
            summary['savings_plans'].append({
                'type': sp.get('savingsPlanType', 'N/A'),
                'commitment': sp.get('commitment', 'N/A'),
                'end_date': sp.get('end', 'N/A'),
                'utilization': sp.get('utilization', {}).get('utilizationPercentage', 'N/A')
            })
    except Exception as e:
        print(f"Savings Plans check: {e}")

    # Authoritative coverage & utilization from Cost Explorer (last 30 days).
    # This is the correct basis - normalized units / hours / spend - versus the
    # naive instance-count ratio above (which ignores RI size-flexibility and
    # does not measure Savings Plans coverage at all). CE is a global endpoint,
    # so it is always called in us-east-1.
    try:
        ce = session.client('ce', region_name='us-east-1')
        end = datetime.now(timezone.utc).date()
        start = end - timedelta(days=30)
        period = {'Start': start.strftime('%Y-%m-%d'), 'End': end.strftime('%Y-%m-%d')}

        try:
            cov = ce.get_reservation_coverage(TimePeriod=period)
            pct = cov.get('Total', {}).get('CoverageHours', {}).get('CoverageHoursPercentage')
            if pct is not None:
                # Keep instance-count `ri_coverage_pct` consistent with the
                # covered/running instance counts shown alongside it; expose the
                # (more accurate) Cost Explorer hours-based coverage separately.
                summary['ri_coverage_hours_pct'] = round(float(pct), 1)
                summary['ri_coverage_basis'] = 'cost_explorer_hours'
        except Exception as e:
            print(f"RI coverage (CE) unavailable, using instance-count estimate: {e}")
            summary.setdefault('ri_coverage_basis', 'instance_count_estimate')

        try:
            sp_cov = ce.get_savings_plans_coverage(TimePeriod=period)
            sp_list = sp_cov.get('SavingsPlansCoverages', [])
            if sp_list:
                sp_pct = sp_list[0].get('Coverage', {}).get('CoveragePercentage')
                if sp_pct is not None:
                    summary['sp_coverage_pct'] = round(float(sp_pct), 1)
        except Exception as e:
            print(f"SP coverage (CE) unavailable: {e}")

        try:
            util = ce.get_savings_plans_utilization(TimePeriod=period)
            up = util.get('Total', {}).get('Utilization', {}).get('UtilizationPercentage')
            if up is not None:
                summary['sp_utilization_pct'] = round(float(up), 1)
        except Exception as e:
            print(f"SP utilization (CE) unavailable: {e}")
    except Exception as e:
        print(f"Cost Explorer coverage lookup skipped: {e}")
        summary.setdefault('ri_coverage_basis', 'instance_count_estimate')

    return summary


def _setup_report_styles(doc):
    """Configure document margins and styles for professional CloudThat reports."""
    for section in doc.sections:
        # Preserve top margin from template for letterhead header
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 82, 147)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 102, 153)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def _add_cover_page(doc, client_name, total_savings, total_recommendations):
    """Add a professional cover page with CloudThat branding."""
    for _ in range(3):
        doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run('CloudThat')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run('\u2501' * 30)
    r.font.color.rgb = RGBColor(0, 102, 153)
    r.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('AWS Infrastructure\nOptimization Report')
    r.font.name = 'Calibri Light'
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_after = Pt(24)

    for label, value in [
        ('Client', client_name),
        ('Date', datetime.now(timezone.utc).strftime('%B %d, %Y')),
        ('Potential Monthly Savings', f'${total_savings:,.2f}'),
        ('Total Recommendations', str(total_recommendations)),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{label}: ')
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(80, 80, 80)
        r = p.add_run(value)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 51, 102)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = conf.add_run('CONFIDENTIAL')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()


def _add_table_of_contents(doc, recommendations, ri_sp_summary):
    """Add a table of contents page."""
    doc.add_heading('Table of Contents', 1)

    sections = ['1. Executive Summary']
    if ri_sp_summary:
        sections.append('2. Reserved Instance & Savings Plans Coverage')

    section_num = 3 if ri_sp_summary else 2
    sections.append(f'{section_num}. Savings Overview Charts')
    section_num += 1

    service_labels = {
        'ec2': 'EC2 Instance Recommendations',
        'stopped_ec2': 'Stopped EC2 Instances',
        'ebs': 'EBS Volume Recommendations',
        'rds': 'RDS Instance Recommendations',
        'lambda': 'Lambda Function Recommendations',
        'eip': 'Elastic IP Recommendations',
        'natgateway': 'NAT Gateway Recommendations',
        's3': 'S3 Bucket Recommendations',
        'dynamodb': 'DynamoDB Table Recommendations',
    }
    for key, label in service_labels.items():
        if key in recommendations and recommendations[key]:
            sections.append(f'{section_num}. {label}')
            section_num += 1

    sections.append(f'{section_num}. Implementation Notes')

    for s in sections:
        p = doc.add_paragraph()
        r = p.add_run(s)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 51, 102)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def _generate_savings_chart(recommendations):
    """Generate a pie chart of savings breakdown by service and return as BytesIO image."""
    service_labels = {
        'ec2': 'EC2', 'stopped_ec2': 'Stopped EC2', 'ebs': 'EBS', 'rds': 'RDS',
        'lambda': 'Lambda', 'eip': 'Elastic IP', 'natgateway': 'NAT Gateway',
        's3': 'S3', 'dynamodb': 'DynamoDB',
    }
    savings_data = {}
    for key, label in service_labels.items():
        if key in recommendations and isinstance(recommendations[key], list):
            total = sum(r.get('monthly_savings', 0) for r in recommendations[key])
            if total > 0:
                savings_data[label] = total

    if not savings_data:
        return None

    colors = ['#003366', '#005293', '#0066CC', '#3399FF', '#66B2FF',
              '#99CCFF', '#CCE5FF', '#006699', '#008080']

    fig, ax = plt.subplots(figsize=(7, 4))
    labels = list(savings_data.keys())
    values = list(savings_data.values())
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, autopct=lambda pct: f'${sum(values)*pct/100:,.0f}\n({pct:.1f}%)',
        colors=colors[:len(values)], startangle=90, textprops={'fontsize': 9}
    )
    for t in autotexts:
        t.set_fontsize(8)
    ax.set_title('Monthly Savings Breakdown by Service', fontsize=13, fontweight='bold',
                 color='#003366', pad=15)
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def _generate_recommendations_chart(recommendations):
    """Generate a bar chart of recommendation counts by service and confidence."""
    service_labels = {
        'ec2': 'EC2', 'stopped_ec2': 'Stopped EC2', 'ebs': 'EBS', 'rds': 'RDS',
        'lambda': 'Lambda', 'eip': 'Elastic IP', 'natgateway': 'NAT GW',
        's3': 'S3', 'dynamodb': 'DynamoDB',
    }
    services = []
    high_counts = []
    medium_counts = []
    for key, label in service_labels.items():
        if key in recommendations and isinstance(recommendations[key], list) and recommendations[key]:
            high = len([r for r in recommendations[key] if r.get('confidence') == 'High'])
            medium = len([r for r in recommendations[key] if r.get('confidence') == 'Medium'])
            services.append(label)
            high_counts.append(high)
            medium_counts.append(medium)

    if not services:
        return None

    fig, ax = plt.subplots(figsize=(7, 4))
    x = range(len(services))
    width = 0.35
    ax.bar([i - width/2 for i in x], high_counts, width, label='High Priority',
           color='#003366', edgecolor='white')
    ax.bar([i + width/2 for i in x], medium_counts, width, label='Medium Priority',
           color='#66B2FF', edgecolor='white')
    ax.set_ylabel('Number of Recommendations', fontsize=10)
    ax.set_title('Recommendations by Service & Priority', fontsize=13,
                 fontweight='bold', color='#003366', pad=15)
    ax.set_xticks(list(x))
    ax.set_xticklabels(services, rotation=30, ha='right', fontsize=9)
    ax.legend(fontsize=9)
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def _generate_savings_by_region_chart(recommendations):
    """Generate a horizontal bar chart of savings by AWS region."""
    region_savings = {}
    for recs in recommendations.values():
        if isinstance(recs, list):
            for r in recs:
                region = r.get('region', 'Unknown')
                region_savings[region] = region_savings.get(region, 0) + r.get('monthly_savings', 0)

    region_savings = {k: v for k, v in region_savings.items() if v > 0}
    # Skip chart for single-region deployments as it provides no comparative value
    if not region_savings or len(region_savings) < 2:
        return None

    sorted_regions = sorted(region_savings.items(), key=lambda x: x[1], reverse=True)
    regions = [r[0] for r in sorted_regions]
    savings = [r[1] for r in sorted_regions]

    fig, ax = plt.subplots(figsize=(7, max(3, len(regions) * 0.5)))
    bars = ax.barh(regions, savings, color='#003366', edgecolor='white', height=0.6)
    for bar, val in zip(bars, savings):
        ax.text(bar.get_width() + max(savings) * 0.02, bar.get_y() + bar.get_height()/2,
                f'${val:,.0f}', va='center', fontsize=9)
    ax.set_xlabel('Monthly Savings ($)', fontsize=10)
    ax.set_title('Monthly Savings by Region', fontsize=13, fontweight='bold',
                 color='#003366', pad=15)
    ax.invert_yaxis()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_styled_table(doc, headers, rows_data, header_color='003366'):
    """Add a professionally styled table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = 1  # CENTER

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, header_color)

    # Add data rows with alternating colors
    for row_idx, row_data in enumerate(rows_data):
        row = table.add_row()
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
            if row_idx % 2 == 1:
                set_cell_background(cell, 'E8F0FE')

    return table


def _add_kpi_table(doc, metrics):
    """Add a key performance indicator table (2-column layout for summary stats)."""
    num_cols = min(len(metrics), 4)
    table = doc.add_table(rows=2, cols=num_cols)
    table.alignment = 1

    for i, (label, value, color) in enumerate(metrics[:num_cols]):
        # Value row
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(value))
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(*color)
        r.font.name = 'Calibri'
        set_cell_background(cell, 'F2F7FC')
        # Label row
        cell = table.rows[1].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(80, 80, 80)
        r.font.name = 'Calibri'
        set_cell_background(cell, 'F2F7FC')

    # Remove table borders for a card look
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'D0D0D0')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    return table


def generate_word_report(recommendations, total_savings, client_name, ri_sp_summary=None, result=None):
    """Generate a professionally formatted Word document report with CloudThat branding."""
    # Load CloudThat template if available, otherwise create blank document
    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()

    _setup_report_styles(doc)

    # Compute summary stats
    total_recommendations = sum(len(v) for v in recommendations.values() if isinstance(v, list))
    high_priority = sum(
        len([r for r in recs if r.get('confidence') == 'High'])
        for recs in recommendations.values() if isinstance(recs, list)
    )
    medium_priority = sum(
        len([r for r in recs if r.get('confidence') == 'Medium'])
        for recs in recommendations.values() if isinstance(recs, list)
    )
    regions_found = set()
    for recs in recommendations.values():
        if isinstance(recs, list):
            for r in recs:
                if 'region' in r:
                    regions_found.add(r['region'])

    # ===== COVER PAGE =====
    _add_cover_page(doc, client_name, total_savings, total_recommendations)

    # ===== TABLE OF CONTENTS =====
    _add_table_of_contents(doc, recommendations, ri_sp_summary)

    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('1. Executive Summary', 1)

    p = doc.add_paragraph()
    r = p.add_run(
        f'This report provides a comprehensive analysis of AWS infrastructure optimization opportunities '
        f'for {client_name}. The analysis covers {len(regions_found)} region(s) and identifies '
        f'{total_recommendations} actionable recommendations across {sum(1 for v in recommendations.values() if isinstance(v, list) and v)} '
        f'AWS services.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(12)

    # KPI cards
    _add_kpi_table(doc, [
        ('Monthly Savings', f'${total_savings:,.2f}', (0, 128, 0)),
        ('Annual Savings', f'${total_savings * 12:,.2f}', (0, 51, 102)),
        ('High Priority', str(high_priority), (192, 0, 0)),
        ('Medium Priority', str(medium_priority), (204, 153, 0)),
    ])

    doc.add_paragraph('')

    if regions_found:
        p = doc.add_paragraph()
        r = p.add_run('Regions Scanned: ')
        r.font.bold = True
        r.font.size = Pt(10)
        r = p.add_run(', '.join(sorted(regions_found)))
        r.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run('Report Generated: ')
    r.font.bold = True
    r.font.size = Pt(10)
    r = p.add_run(datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC'))
    r.font.size = Pt(10)

    # Forecast + rate-optimization context (when available from Cost Explorer).
    forecast = (result or {}).get('forecast') or {}
    commitment_savings = (result or {}).get('commitment_savings', 0) or 0
    if forecast.get('forecast_month') is not None or commitment_savings > 0:
        p = doc.add_paragraph()
        if forecast.get('forecast_month') is not None:
            r = p.add_run(f"Forecasted spend (this month): ${forecast['forecast_month']:,.2f}   ")
            r.font.size = Pt(10)
        if forecast.get('month_to_date') is not None:
            r = p.add_run(f"Month-to-date: ${forecast['month_to_date']:,.2f}   ")
            r.font.size = Pt(10)
        if commitment_savings > 0:
            r = p.add_run(f"Additional Savings Plans opportunity: ${commitment_savings:,.2f}/mo")
            r.font.size = Pt(10)
            r.font.bold = True

    # Top quick wins (low effort + low risk) - the actions to do first.
    all_recs = []
    for svc, recs in recommendations.items():
        if isinstance(recs, list):
            for rr in recs:
                if rr.get('monthly_savings', 0) > 0:
                    all_recs.append((svc, rr))
    quick = sorted([x for x in all_recs if x[1].get('quick_win')],
                   key=lambda x: x[1].get('monthly_savings', 0), reverse=True)[:5]
    if quick:
        doc.add_heading('Top Quick Wins (Low Effort, Low Risk)', 2)
        qrows = []
        for svc, rr in quick:
            res_id = (rr.get('instance_id') or rr.get('volume_id') or rr.get('db_id')
                      or rr.get('function_name') or rr.get('snapshot_id')
                      or rr.get('load_balancer_name') or rr.get('table_name')
                      or rr.get('ip_address') or rr.get('bucket_name') or 'resource')
            qrows.append([SERVICE_LABELS.get(svc, svc), str(res_id),
                          f"${rr.get('monthly_savings', 0):,.2f}",
                          f"${rr.get('annual_savings', 0):,.2f}"])
        _add_styled_table(doc, ['Service', 'Resource', 'Monthly', 'Annual'], qrows)
        doc.add_paragraph('')

    # Service-level savings summary table
    doc.add_heading('Savings Summary by Service', 2)
    summary_headers = ['Service', 'Recommendations', 'Monthly Savings', 'Annual Savings']
    summary_rows = []
    service_order = [
        ('ec2', 'EC2 Instances'), ('stopped_ec2', 'Stopped EC2'), ('ebs', 'EBS Volumes'),
        ('ebs_snapshot', 'EBS Snapshots'), ('rds', 'RDS Instances'),
        ('lambda', 'Lambda Functions'), ('eip', 'Elastic IPs / Public IPv4'),
        ('elb', 'Load Balancers'), ('natgateway', 'NAT Gateways'),
        ('s3', 'S3 Buckets'), ('dynamodb', 'DynamoDB Tables'),
    ]
    for key, label in service_order:
        if key in recommendations and isinstance(recommendations[key], list) and recommendations[key]:
            svc_savings = sum(r.get('monthly_savings', 0) for r in recommendations[key])
            summary_rows.append([
                label, str(len(recommendations[key])),
                f'${svc_savings:,.2f}', f'${svc_savings * 12:,.2f}'
            ])
    if summary_rows:
        summary_rows.append(['TOTAL', str(total_recommendations),
                             f'${total_savings:,.2f}', f'${total_savings * 12:,.2f}'])
        _add_styled_table(doc, summary_headers, summary_rows)

    doc.add_page_break()

    # ===== RI/SP COVERAGE =====
    if ri_sp_summary:
        doc.add_heading('2. Reserved Instance & Savings Plans Coverage', 1)

        coverage_pct = ri_sp_summary.get('ri_coverage_pct', 0)
        total_instances = ri_sp_summary.get('total_running_instances', 0)
        ri_covered = ri_sp_summary.get('ri_covered_instances', 0)

        _add_kpi_table(doc, [
            ('RI Coverage', f'{coverage_pct:.1f}%',
             (0, 128, 0) if coverage_pct >= 50 else (192, 0, 0)),
            ('Running Instances', str(total_instances), (0, 51, 102)),
            ('RI-Covered', str(ri_covered), (0, 82, 147)),
            ('Uncovered', str(total_instances - ri_covered), (204, 153, 0)),
        ])

        doc.add_paragraph('')

        if ri_sp_summary.get('active_ris'):
            doc.add_heading('Active Reserved Instances', 2)
            ri_rows = [[ri['instance_type'], str(ri['count']), ri['offering_type'], str(ri['end_date'])]
                       for ri in ri_sp_summary['active_ris']]
            _add_styled_table(doc, ['Instance Type', 'Count', 'Offering Type', 'End Date'], ri_rows)

        doc.add_paragraph('')

        if ri_sp_summary.get('savings_plans'):
            doc.add_heading('Active Savings Plans', 2)
            sp_rows = [[str(sp['type']), str(sp['commitment']), str(sp['end_date'])]
                       for sp in ri_sp_summary['savings_plans']]
            _add_styled_table(doc, ['Type', 'Commitment', 'End Date'], sp_rows)

        doc.add_page_break()

    # ===== CHARTS =====
    chart_section_num = 3 if ri_sp_summary else 2
    doc.add_heading(f'{chart_section_num}. Savings Overview', 1)

    savings_chart = _generate_savings_chart(recommendations)
    if savings_chart:
        doc.add_picture(savings_chart, width=Inches(5.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

    recs_chart = _generate_recommendations_chart(recommendations)
    if recs_chart:
        doc.add_picture(recs_chart, width=Inches(5.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

    region_chart = _generate_savings_by_region_chart(recommendations)
    if region_chart:
        doc.add_picture(region_chart, width=Inches(5.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== SERVICE RECOMMENDATIONS =====
    section_num = chart_section_num + 1

    # EC2 Instances
    if 'ec2' in recommendations and recommendations['ec2']:
        doc.add_heading(f'{section_num}. EC2 Instance Recommendations', 1)
        section_num += 1
        ec2_rows = []
        for rec in recommendations['ec2']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            basis_note = ' [savings after RI/SP discount]' if rec.get('savings_basis') == 'after_discounts' else ''
            ec2_rows.append([
                f"{region_prefix}{rec['instance_id']}",
                format_tags_str(rec.get('tags', {})),
                rec['current_type'], rec['recommended_type'],
                f"${rec['current_cost']:.2f}", f"${rec['recommended_cost']:.2f}",
                f"${rec['monthly_savings']:.2f}",
                f"{rec['reason']} (CPU: {rec['cpu_avg']}%){basis_note}",
                rec['confidence'],
            ])
        table = _add_styled_table(doc, [
            'Instance ID', 'Tags', 'Current Type', 'Recommended',
            'Current Cost', 'New Cost', 'Monthly Savings', 'Reason', 'Confidence'
        ], ec2_rows)
        # Color-code High confidence cells
        for row_idx, rec in enumerate(recommendations['ec2']):
            if rec['confidence'] == 'High':
                set_cell_background(table.rows[row_idx + 1].cells[8], '90EE90')
        doc.add_paragraph('')

    # Stopped EC2 Instances
    if 'stopped_ec2' in recommendations and recommendations['stopped_ec2']:
        doc.add_heading(f'{section_num}. Stopped EC2 Instances (EBS Cost Waste)', 1)
        section_num += 1
        rows = []
        for rec in recommendations['stopped_ec2']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['instance_id']}",
                format_tags_str(rec.get('tags', {})),
                rec['instance_type'], str(rec['stopped_days']),
                str(rec['attached_volumes']),
                f"${rec['monthly_savings']:.2f}", rec['recommendation'],
            ])
        _add_styled_table(doc, [
            'Instance ID', 'Tags', 'Type', 'Days Stopped',
            'Attached Volumes', 'Monthly EBS Cost', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # EBS Volumes
    if 'ebs' in recommendations and recommendations['ebs']:
        doc.add_heading(f'{section_num}. EBS Volume Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['ebs']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['volume_id']}",
                format_tags_str(rec.get('tags', {})),
                str(rec['size']), rec['type'], rec['issue'],
                rec['recommendation'], f"${rec['monthly_savings']:.2f}",
            ])
        _add_styled_table(doc, [
            'Volume ID', 'Tags', 'Size (GB)', 'Type', 'Issue',
            'Recommendation', 'Monthly Savings'
        ], rows)
        doc.add_paragraph('')

    # RDS Instances
    if 'rds' in recommendations and recommendations['rds']:
        doc.add_heading(f'{section_num}. RDS Instance Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['rds']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['db_id']}",
                format_tags_str(rec.get('tags', {})),
                rec['current_class'], rec['recommended_class'],
                f"${rec['current_cost']:.2f}", f"${rec['recommended_cost']:.2f}",
                f"${rec['monthly_savings']:.2f}", rec['reason'],
            ])
        _add_styled_table(doc, [
            'DB Identifier', 'Tags', 'Current Class', 'Recommended',
            'Current Cost', 'New Cost', 'Monthly Savings', 'Reason'
        ], rows)
        doc.add_paragraph('')

    # Lambda Functions
    if 'lambda' in recommendations and recommendations['lambda']:
        doc.add_heading(f'{section_num}. Lambda Function Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['lambda']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['function_name']}",
                format_tags_str(rec.get('tags', {})),
                f"{rec['current_memory']} MB", f"{rec['recommended_memory']} MB",
                str(int(rec['avg_duration'])),
                f"${rec['current_cost']:.2f}", f"${rec['recommended_cost']:.2f}",
                f"${rec['monthly_savings']:.2f}",
            ])
        _add_styled_table(doc, [
            'Function Name', 'Tags', 'Current Memory', 'Recommended',
            'Avg Duration (ms)', 'Current Cost', 'New Cost', 'Monthly Savings'
        ], rows)
        doc.add_paragraph('')

    # Elastic IPs
    if 'eip' in recommendations and recommendations['eip']:
        doc.add_heading(f'{section_num}. Elastic IP Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['eip']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['ip_address']}",
                format_tags_str(rec.get('tags', {})),
                rec['status'], f"${rec['monthly_savings']:.2f}",
                rec['recommendation'],
            ])
        _add_styled_table(doc, [
            'IP Address', 'Tags', 'Status', 'Monthly Cost', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # NAT Gateways
    if 'natgateway' in recommendations and recommendations['natgateway']:
        doc.add_heading(f'{section_num}. NAT Gateway Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['natgateway']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['nat_gateway_id']}",
                format_tags_str(rec.get('tags', {})),
                rec.get('vpc_id', 'N/A'), f"{rec['avg_daily_gb']:.2f}",
                f"${rec['monthly_cost']:.2f}", rec.get('reason', ''),
                rec['recommendation'],
            ])
        _add_styled_table(doc, [
            'NAT Gateway ID', 'Tags', 'VPC', 'Avg Daily GB',
            'Monthly Cost', 'Reason', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # S3 Buckets
    if 's3' in recommendations and recommendations['s3']:
        doc.add_heading(f'{section_num}. S3 Bucket Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['s3']:
            rows.append([
                rec['bucket_name'], format_tags_str(rec.get('tags', {})),
                rec.get('region', 'N/A'), rec['issues'], rec['recommendation'],
            ])
        _add_styled_table(doc, [
            'Bucket Name', 'Tags', 'Region', 'Issues', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # DynamoDB Tables
    if 'dynamodb' in recommendations and recommendations['dynamodb']:
        doc.add_heading(f'{section_num}. DynamoDB Table Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['dynamodb']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['table_name']}",
                format_tags_str(rec.get('tags', {})),
                f"{rec['provisioned_rcu']}/{rec['provisioned_wcu']}",
                f"{rec['avg_rcu']}/{rec['avg_wcu']}",
                f"RCU: {rec['rcu_utilization']}%, WCU: {rec['wcu_utilization']}%",
                f"${rec['current_cost']:.2f}", f"${rec['monthly_savings']:.2f}",
                rec['recommendation'],
            ])
        _add_styled_table(doc, [
            'Table Name', 'Tags', 'Provisioned RCU/WCU', 'Avg RCU/WCU',
            'Utilization', 'Current Cost', 'Monthly Savings', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # EBS Snapshots
    if 'ebs_snapshot' in recommendations and recommendations['ebs_snapshot']:
        doc.add_heading(f'{section_num}. EBS Snapshot Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['ebs_snapshot']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['snapshot_id']}",
                format_tags_str(rec.get('tags', {})),
                str(rec.get('size', 'N/A')), str(rec.get('age_days', 'N/A')),
                rec.get('issue', ''), f"${rec['monthly_savings']:.2f}",
                rec.get('recommendation', ''),
            ])
        _add_styled_table(doc, [
            'Snapshot ID', 'Tags', 'Size (GB)', 'Age (days)', 'Issue',
            'Monthly Savings', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # Load Balancers
    if 'elb' in recommendations and recommendations['elb']:
        doc.add_heading(f'{section_num}. Idle Load Balancer Recommendations', 1)
        section_num += 1
        rows = []
        for rec in recommendations['elb']:
            region_prefix = f"[{rec.get('region', '')}] " if rec.get('region') else ''
            rows.append([
                f"{region_prefix}{rec['load_balancer_name']}",
                format_tags_str(rec.get('tags', {})),
                rec.get('type', ''), rec.get('metric', ''),
                f"${rec['monthly_savings']:.2f}", rec.get('recommendation', ''),
            ])
        _add_styled_table(doc, [
            'Load Balancer', 'Tags', 'Type', 'Traffic', 'Monthly Savings', 'Recommendation'
        ], rows)
        doc.add_paragraph('')

    # Savings Plans purchase (rate optimization - shown separately)
    if 'savings_plan' in recommendations and recommendations['savings_plan']:
        doc.add_heading(f'{section_num}. Savings Plans Purchase (Rate Optimization)', 1)
        section_num += 1
        p = doc.add_paragraph()
        r = p.add_run('These savings come from purchasing a commitment, not from '
                      'eliminating waste, and are reported separately. Rightsize idle/'
                      'over-provisioned resources first, then commit to the steady-state baseline.')
        r.font.size = Pt(10)
        r.font.italic = True
        rows = []
        for rec in recommendations['savings_plan']:
            rows.append([
                rec.get('type', 'Compute Savings Plan'), rec.get('term', ''),
                rec.get('payment_option', ''), str(rec.get('hourly_commitment', 'N/A')),
                f"{rec.get('estimated_savings_pct', 'N/A')}%",
                f"${rec['monthly_savings']:.2f}",
            ])
        _add_styled_table(doc, [
            'Type', 'Term', 'Payment', 'Hourly Commit', 'Est. Savings %', 'Monthly Savings'
        ], rows)
        doc.add_paragraph('')

    # ===== IMPLEMENTATION NOTES =====
    doc.add_page_break()
    doc.add_heading(f'{section_num}. Implementation Notes & Methodology', 1)

    # Methodology / basis of numbers (transparency builds trust).
    p = doc.add_paragraph()
    r = p.add_run('Methodology & Basis: ')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0, 51, 102)
    r = p.add_run(
        'Recommendations combine AWS Compute Optimizer (ML, ~14-day metrics), '
        'CloudWatch utilization, and live AWS Price List pricing. Savings are '
        'On-Demand-based unless a recommendation is marked "after discounts". '
        'Waste-elimination savings and Savings Plans (rate) savings are reported '
        'separately and are not summed, to avoid double counting.'
    )
    r.font.size = Pt(10)

    notes = [
        ('High Confidence', 'Based on AWS Compute Optimizer ML analysis with robust data points. '
         'These recommendations carry the highest level of certainty.', (0, 128, 0)),
        ('Medium Confidence', 'Based on 14-day CloudWatch metrics analysis. '
         'Review current workload patterns before implementing.', (204, 153, 0)),
    ]
    for title_text, desc, color in notes:
        p = doc.add_paragraph()
        r = p.add_run(f'{title_text}: ')
        r.font.bold = True
        r.font.color.rgb = RGBColor(*color)
        r.font.size = Pt(11)
        r = p.add_run(desc)
        r.font.size = Pt(10)

    doc.add_heading('Best Practices', 2)
    best_practices = [
        'Test all changes in non-production environments before applying to production.',
        'Implement changes during maintenance windows to minimize disruption.',
        'Consider Reserved Instances and Savings Plans before rightsizing.',
        'Monitor performance metrics closely for 48-72 hours after implementing changes.',
        'Review and update recommendations quarterly as workload patterns evolve.',
        'Use AWS Cost Explorer to track actual savings after implementation.',
    ]
    for practice in best_practices:
        p = doc.add_paragraph()
        r = p.add_run(f'\u2022  {practice}')
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

    # ===== DISCLAIMER =====
    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run('Disclaimer: ')
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)
    r = p.add_run(
        'This report is generated based on automated analysis of AWS CloudWatch metrics and '
        'Compute Optimizer data over a 14-day period. Recommendations should be validated '
        'against business requirements and application-specific constraints before implementation. '
        'CloudThat is not liable for any disruption caused by implementing these recommendations '
        'without proper testing.'
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)

    return doc

def set_cell_background(cell, color):
    """Set cell background color (hex color without #)"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)

def get_metric_value(rec, metric_name):
    """Extract metric value from Compute Optimizer recommendation"""
    for metric in rec.get('utilizationMetrics', []):
        if metric['name'] == metric_name:
            return round(metric['value'], 1)
    return 'N/A'


def get_resource_tags(tags_list):
    """Extract Name, Owner, and Team tags from a tags list."""
    result = {}
    for tag in (tags_list or []):
        key = tag.get('Key', '')
        if key == 'Name':
            result['Name'] = tag.get('Value', '')
        elif key.lower() == 'owner':
            result['Owner'] = tag.get('Value', '')
        elif key.lower() == 'team':
            result['Team'] = tag.get('Value', '')
    return result


def format_tags_str(tags):
    """Format tags dict into a readable string."""
    if not tags:
        return 'N/A'
    parts = []
    for key in ['Name', 'Owner', 'Team']:
        if key in tags:
            parts.append(f"{key}: {tags[key]}")
    return ', '.join(parts) if parts else 'N/A'


class PricingUnavailableError(Exception):
    """Raised when pricing data cannot be fetched from AWS Pricing API"""
    pass


def get_instance_cost(instance_type, region, operating_system='Linux',
                      tenancy='Shared', pre_installed_sw='NA',
                      license_model='No License required'):
    """Get actual EC2 on-demand hourly price from the AWS Price List API (cached).

    Pricing is OS/tenancy/software aware: a Windows or SQL-Server box can cost
    ~2x a Linux box of the same shape, so pricing everything as Linux/Shared
    (the previous behaviour) made every derived saving wrong for those instances.

    Raises PricingUnavailableError if pricing cannot be fetched - no fallback to
    ensure we never fabricate numbers.
    """
    cache_key = f"ec2_{instance_type}_{region}_{operating_system}_{tenancy}_{pre_installed_sw}"

    # Check cache first
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']

    try:
        pricing_client = boto3.client('pricing', region_name='us-east-1')

        location = REGION_LOCATION_MAP.get(region)
        if not location:
            raise PricingUnavailableError(f"Unknown region: {region}")

        filters = [
            {'Type': 'TERM_MATCH', 'Field': 'instanceType', 'Value': instance_type},
            {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
            {'Type': 'TERM_MATCH', 'Field': 'operatingSystem', 'Value': operating_system},
            {'Type': 'TERM_MATCH', 'Field': 'tenancy', 'Value': tenancy},
            {'Type': 'TERM_MATCH', 'Field': 'preInstalledSw', 'Value': pre_installed_sw},
            {'Type': 'TERM_MATCH', 'Field': 'capacitystatus', 'Value': 'Used'},
        ]
        # licenseModel only applies to licensed platforms (Windows/RHEL/SUSE).
        if operating_system != 'Linux':
            filters.append({'Type': 'TERM_MATCH', 'Field': 'licenseModel', 'Value': license_model})

        response = pricing_client.get_products(
            ServiceCode='AmazonEC2',
            Filters=filters,
            MaxResults=10
        )
        
        if response['PriceList']:
            # Parse all results and find the on-demand price
            for price_item in response['PriceList']:
                price_data = json.loads(price_item)
                on_demand = price_data.get('terms', {}).get('OnDemand', {})
                if on_demand:
                    price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                    for dim in price_dimensions:
                        price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                        if price_per_unit and float(price_per_unit) > 0:
                            price_per_hour = float(price_per_unit)
                            
                            # Cache the result
                            PRICING_CACHE[cache_key] = {
                                'price': price_per_hour,
                                'timestamp': datetime.now().timestamp()
                            }
                            return price_per_hour
        
        # No pricing found - raise error instead of fallback
        raise PricingUnavailableError(f"No pricing found for EC2 instance type {instance_type} in {region}")
        
    except PricingUnavailableError:
        raise
    except Exception as e:
        raise PricingUnavailableError(f"Failed to fetch EC2 pricing for {instance_type} in {region}: {e}")

def get_rds_cost(db_class, engine, region, multi_az=False):
    """Get actual RDS pricing with caching.
    
    Raises PricingUnavailableError if pricing cannot be fetched - no fallback to ensure accuracy.
    """
    cache_key = f"rds_{db_class}_{engine}_{region}_{'multiaz' if multi_az else 'singleaz'}"
    
    # Check cache first
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']
    
    try:
        pricing_client = boto3.client('pricing', region_name='us-east-1')
        
        location = REGION_LOCATION_MAP.get(region)
        if not location:
            raise PricingUnavailableError(f"Unknown region: {region}")
        
        # Map engine names to pricing API values
        engine_map = {
            'postgres': 'PostgreSQL', 'mysql': 'MySQL', 'mariadb': 'MariaDB',
            'oracle-se': 'Oracle', 'oracle-se1': 'Oracle', 'oracle-se2': 'Oracle', 'oracle-ee': 'Oracle',
            'sqlserver-se': 'SQL Server', 'sqlserver-ee': 'SQL Server', 'sqlserver-ex': 'SQL Server', 'sqlserver-web': 'SQL Server',
            'aurora': 'Aurora MySQL', 'aurora-mysql': 'Aurora MySQL', 'aurora-postgresql': 'Aurora PostgreSQL'
        }
        db_engine = engine_map.get(engine.lower())
        if not db_engine:
            raise PricingUnavailableError(f"Unknown RDS engine: {engine}")
        
        deployment_option = 'Multi-AZ' if multi_az else 'Single-AZ'
        
        response = pricing_client.get_products(
            ServiceCode='AmazonRDS',
            Filters=[
                {'Type': 'TERM_MATCH', 'Field': 'instanceType', 'Value': db_class},
                {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                {'Type': 'TERM_MATCH', 'Field': 'databaseEngine', 'Value': db_engine},
                {'Type': 'TERM_MATCH', 'Field': 'deploymentOption', 'Value': deployment_option}
            ],
            MaxResults=10
        )
        
        if response['PriceList']:
            for price_item in response['PriceList']:
                price_data = json.loads(price_item)
                on_demand = price_data.get('terms', {}).get('OnDemand', {})
                if on_demand:
                    price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                    for dim in price_dimensions:
                        price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                        if price_per_unit and float(price_per_unit) > 0:
                            price_per_hour = float(price_per_unit)
                            
                            # Cache the result
                            PRICING_CACHE[cache_key] = {
                                'price': price_per_hour,
                                'timestamp': datetime.now().timestamp()
                            }
                            return price_per_hour
        
        # No pricing found - raise error instead of fallback
        raise PricingUnavailableError(f"No pricing found for RDS {db_class} ({engine}) in {region}")
        
    except PricingUnavailableError:
        raise
    except Exception as e:
        raise PricingUnavailableError(f"Failed to fetch RDS pricing for {db_class} in {region}: {e}")

def calculate_ebs_cost(volume_type, size_gb, region, iops=0, throughput=0):
    """Calculate monthly EBS cost with real-time pricing.
    
    Raises PricingUnavailableError if pricing cannot be fetched - no fallback to ensure accuracy.
    """
    cache_key = f"ebs_{volume_type}_{region}"
    
    # Check cache for base price
    base_price_per_gb = None
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            base_price_per_gb = cached_data['price']
    
    if base_price_per_gb is None:
        try:
            pricing_client = boto3.client('pricing', region_name='us-east-1')
            location = REGION_LOCATION_MAP.get(region)
            if not location:
                raise PricingUnavailableError(f"Unknown region: {region}")
            
            response = pricing_client.get_products(
                ServiceCode='AmazonEC2',
                Filters=[
                    {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'Storage'},
                    {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                    {'Type': 'TERM_MATCH', 'Field': 'volumeApiName', 'Value': volume_type}
                ],
                MaxResults=10
            )
            
            if response['PriceList']:
                for price_item in response['PriceList']:
                    price_data = json.loads(price_item)
                    on_demand = price_data.get('terms', {}).get('OnDemand', {})
                    if on_demand:
                        price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                        for dim in price_dimensions:
                            # Look for per GB-month pricing
                            if 'GB-Mo' in dim.get('unit', ''):
                                price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                                if price_per_unit:
                                    base_price_per_gb = float(price_per_unit)
                                    # Cache it
                                    PRICING_CACHE[cache_key] = {
                                        'price': base_price_per_gb,
                                        'timestamp': datetime.now().timestamp()
                                    }
                                    break
            
            if base_price_per_gb is None:
                raise PricingUnavailableError(f"No pricing found for EBS volume type {volume_type} in {region}")
                
        except PricingUnavailableError:
            raise
        except Exception as e:
            raise PricingUnavailableError(f"Failed to fetch EBS pricing for {volume_type} in {region}: {e}")
    
    # Calculate total cost
    total_cost = base_price_per_gb * size_gb
    
    # Add IOPS cost for provisioned IOPS volumes - fetch from API
    if volume_type in ['io1', 'io2'] and iops > 0:
        iops_price = get_ebs_iops_cost(volume_type, region)
        total_cost += iops * iops_price
    
    # gp3 additional IOPS/throughput costs (beyond baseline)
    if volume_type == 'gp3':
        # gp3 baseline: 3000 IOPS, 125 MB/s throughput
        if iops > 3000:
            extra_iops = iops - 3000
            gp3_iops_price = get_ebs_gp3_iops_cost(region)
            total_cost += extra_iops * gp3_iops_price
        if throughput > 125:
            extra_throughput = throughput - 125
            gp3_throughput_price = get_ebs_gp3_throughput_cost(region)
            total_cost += extra_throughput * gp3_throughput_price
    
    return total_cost


def get_ebs_iops_cost(volume_type, region):
    """Get EBS IOPS pricing for io1/io2 volumes."""
    cache_key = f"ebs_iops_{volume_type}_{region}"
    
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']
    
    try:
        pricing_client = boto3.client('pricing', region_name='us-east-1')
        location = REGION_LOCATION_MAP.get(region)
        if not location:
            raise PricingUnavailableError(f"Unknown region: {region}")
        
        response = pricing_client.get_products(
            ServiceCode='AmazonEC2',
            Filters=[
                {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'System Operation'},
                {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                {'Type': 'TERM_MATCH', 'Field': 'volumeApiName', 'Value': volume_type}
            ],
            MaxResults=10
        )
        
        if response['PriceList']:
            for price_item in response['PriceList']:
                price_data = json.loads(price_item)
                on_demand = price_data.get('terms', {}).get('OnDemand', {})
                if on_demand:
                    price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                    for dim in price_dimensions:
                        if 'IOPS-Mo' in dim.get('unit', ''):
                            price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                            if price_per_unit:
                                iops_price = float(price_per_unit)
                                PRICING_CACHE[cache_key] = {
                                    'price': iops_price,
                                    'timestamp': datetime.now().timestamp()
                                }
                                return iops_price
        
        raise PricingUnavailableError(f"No IOPS pricing found for {volume_type} in {region}")
        
    except PricingUnavailableError:
        raise
    except Exception as e:
        raise PricingUnavailableError(f"Failed to fetch EBS IOPS pricing for {volume_type} in {region}: {e}")


def get_ebs_gp3_iops_cost(region):
    """Get gp3 additional IOPS cost."""
    cache_key = f"ebs_gp3_iops_{region}"
    
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']
    
    try:
        pricing_client = boto3.client('pricing', region_name='us-east-1')
        location = REGION_LOCATION_MAP.get(region)
        if not location:
            raise PricingUnavailableError(f"Unknown region: {region}")
        
        response = pricing_client.get_products(
            ServiceCode='AmazonEC2',
            Filters=[
                {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'System Operation'},
                {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                {'Type': 'TERM_MATCH', 'Field': 'volumeApiName', 'Value': 'gp3'},
                {'Type': 'TERM_MATCH', 'Field': 'group', 'Value': 'EBS IOPS'}
            ],
            MaxResults=10
        )
        
        if response['PriceList']:
            for price_item in response['PriceList']:
                price_data = json.loads(price_item)
                on_demand = price_data.get('terms', {}).get('OnDemand', {})
                if on_demand:
                    price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                    for dim in price_dimensions:
                        price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                        if price_per_unit and float(price_per_unit) > 0:
                            iops_price = float(price_per_unit)
                            PRICING_CACHE[cache_key] = {
                                'price': iops_price,
                                'timestamp': datetime.now().timestamp()
                            }
                            return iops_price
        
        raise PricingUnavailableError(f"No gp3 IOPS pricing found in {region}")
        
    except PricingUnavailableError:
        raise
    except Exception as e:
        raise PricingUnavailableError(f"Failed to fetch gp3 IOPS pricing in {region}: {e}")


def get_ebs_gp3_throughput_cost(region):
    """Get gp3 additional throughput cost."""
    cache_key = f"ebs_gp3_throughput_{region}"
    
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']
    
    try:
        pricing_client = boto3.client('pricing', region_name='us-east-1')
        location = REGION_LOCATION_MAP.get(region)
        if not location:
            raise PricingUnavailableError(f"Unknown region: {region}")
        
        response = pricing_client.get_products(
            ServiceCode='AmazonEC2',
            Filters=[
                {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'System Operation'},
                {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                {'Type': 'TERM_MATCH', 'Field': 'volumeApiName', 'Value': 'gp3'},
                {'Type': 'TERM_MATCH', 'Field': 'group', 'Value': 'EBS Throughput'}
            ],
            MaxResults=10
        )
        
        if response['PriceList']:
            for price_item in response['PriceList']:
                price_data = json.loads(price_item)
                on_demand = price_data.get('terms', {}).get('OnDemand', {})
                if on_demand:
                    price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                    for dim in price_dimensions:
                        price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                        if price_per_unit and float(price_per_unit) > 0:
                            throughput_price = float(price_per_unit)
                            PRICING_CACHE[cache_key] = {
                                'price': throughput_price,
                                'timestamp': datetime.now().timestamp()
                            }
                            return throughput_price
        
        raise PricingUnavailableError(f"No gp3 throughput pricing found in {region}")
        
    except PricingUnavailableError:
        raise
    except Exception as e:
        raise PricingUnavailableError(f"Failed to fetch gp3 throughput pricing in {region}: {e}")


def calculate_lambda_cost(memory_mb, avg_duration_ms, invocations, region='us-east-1'):
    """Calculate monthly Lambda cost with real-time pricing.
    
    Raises PricingUnavailableError if pricing cannot be fetched - no fallback to ensure accuracy.
    """
    cache_key = f"lambda_{region}"
    
    # Try to get real-time pricing
    cost_per_gb_second = None
    request_cost_per_million = None
    
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            cost_per_gb_second = cached_data.get('gb_second')
            request_cost_per_million = cached_data.get('request')
    
    if cost_per_gb_second is None or request_cost_per_million is None:
        try:
            pricing_client = boto3.client('pricing', region_name='us-east-1')
            location = REGION_LOCATION_MAP.get(region)
            if not location:
                raise PricingUnavailableError(f"Unknown region: {region}")
            
            response = pricing_client.get_products(
                ServiceCode='AWSLambda',
                Filters=[
                    {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location}
                ],
                MaxResults=20
            )
            
            if response['PriceList']:
                for price_item in response['PriceList']:
                    price_data = json.loads(price_item)
                    product = price_data.get('product', {})
                    attributes = product.get('attributes', {})
                    group = attributes.get('group', '')
                    
                    on_demand = price_data.get('terms', {}).get('OnDemand', {})
                    if on_demand:
                        price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                        for dim in price_dimensions:
                            price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                            if price_per_unit and float(price_per_unit) > 0:
                                unit = dim.get('unit', '')
                                if 'second' in unit.lower() or 'Lambda-GB-Second' in group:
                                    cost_per_gb_second = float(price_per_unit)
                                elif 'request' in unit.lower() or 'Request' in group:
                                    request_cost_per_million = float(price_per_unit) * 1000000
                
                if cost_per_gb_second and request_cost_per_million:
                    PRICING_CACHE[cache_key] = {
                        'gb_second': cost_per_gb_second,
                        'request': request_cost_per_million,
                        'timestamp': datetime.now().timestamp()
                    }
            
            if cost_per_gb_second is None or request_cost_per_million is None:
                raise PricingUnavailableError(f"No Lambda pricing found in {region}")
                
        except PricingUnavailableError:
            raise
        except Exception as e:
            raise PricingUnavailableError(f"Failed to fetch Lambda pricing in {region}: {e}")
    
    # Calculate costs
    gb_seconds = (memory_mb / 1024) * (avg_duration_ms / 1000) * invocations
    compute_cost = gb_seconds * cost_per_gb_second
    request_cost = invocations * (request_cost_per_million / 1000000)
    
    return compute_cost + request_cost


def get_dynamodb_pricing(region):
    """Get DynamoDB capacity pricing from the Price List API (cached).

    Returns a dict with provisioned RCU/WCU per-hour prices and on-demand
    read/write request-unit prices per million. Falls back to current us-east-1
    published rates (on-demand throughput was cut ~50% on 2024-11-01) if the
    Price List cannot be parsed - so a DynamoDB recommendation is never dropped,
    while remaining far more accurate than the previous stale hardcoded values.
    """
    cache_key = f"dynamodb_pricing_{region}"
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']

    pricing = {
        'rcu_hour': 0.00013, 'wcu_hour': 0.00065,
        'rru_per_million': 0.125, 'wru_per_million': 0.625,
        'source': 'default_us_east_1',
    }

    try:
        location = REGION_LOCATION_MAP.get(region)
        if location:
            pricing_client = boto3.client('pricing', region_name='us-east-1')
            found = {}
            paginator = pricing_client.get_paginator('get_products')
            for page in paginator.paginate(
                ServiceCode='AmazonDynamoDB',
                Filters=[{'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location}],
                PaginationConfig={'MaxItems': 400}
            ):
                for price_item in page['PriceList']:
                    data = json.loads(price_item)
                    on_demand = data.get('terms', {}).get('OnDemand', {})
                    if not on_demand:
                        continue
                    for dim in list(list(on_demand.values())[0]['priceDimensions'].values()):
                        unit = dim.get('unit', '')
                        usd = dim.get('pricePerUnit', {}).get('USD')
                        if not usd:
                            continue
                        val = float(usd)
                        if 'ReadCapacityUnit-Hrs' in unit and val > 0:
                            found['rcu_hour'] = val
                        elif 'WriteCapacityUnit-Hrs' in unit and val > 0:
                            found['wcu_hour'] = val
                        elif 'ReadRequestUnits' in unit and val > 0:
                            found['rru_per_million'] = val * 1_000_000
                        elif 'WriteRequestUnits' in unit and val > 0:
                            found['wru_per_million'] = val * 1_000_000
            # Only trust the Price List result if we recovered the core rates.
            if 'rru_per_million' in found and 'wru_per_million' in found:
                pricing.update(found)
                pricing['source'] = 'price_list'
    except Exception as e:
        print(f"DynamoDB pricing lookup fell back to defaults: {e}")

    PRICING_CACHE[cache_key] = {'price': pricing, 'timestamp': datetime.now().timestamp()}
    return pricing


def get_nat_gateway_pricing(region):
    """Get NAT Gateway hourly + per-GB data-processing price from Price List (cached).

    Falls back to common us-east-1 rates if the Price List cannot be parsed, so a
    recommendation is never dropped. Previously both were hardcoded at $0.045.
    """
    cache_key = f"natgw_pricing_{region}"
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']

    pricing = {'hourly': 0.045, 'per_gb': 0.045, 'source': 'default_us_east_1'}
    try:
        location = REGION_LOCATION_MAP.get(region)
        if location:
            pricing_client = boto3.client('pricing', region_name='us-east-1')
            response = pricing_client.get_products(
                ServiceCode='AmazonEC2',
                Filters=[
                    {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'NAT Gateway'},
                    {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                ],
                MaxResults=20
            )
            found = {}
            for price_item in response.get('PriceList', []):
                data = json.loads(price_item)
                on_demand = data.get('terms', {}).get('OnDemand', {})
                if not on_demand:
                    continue
                for dim in list(list(on_demand.values())[0]['priceDimensions'].values()):
                    unit = dim.get('unit', '').lower()
                    usd = dim.get('pricePerUnit', {}).get('USD')
                    if not usd:
                        continue
                    val = float(usd)
                    if val <= 0:
                        continue
                    if 'hrs' in unit or 'hour' in unit:
                        found['hourly'] = val
                    elif 'gb' in unit:
                        found['per_gb'] = val
            if 'hourly' in found:
                pricing.update(found)
                pricing['source'] = 'price_list'
    except Exception as e:
        print(f"NAT gateway pricing lookup fell back to defaults: {e}")

    PRICING_CACHE[cache_key] = {'price': pricing, 'timestamp': datetime.now().timestamp()}
    return pricing


def get_eip_cost(region):
    """Get Elastic IP / public IPv4 hourly cost from the pricing API.

    Raises PricingUnavailableError if pricing cannot be fetched - no fallback to
    ensure accuracy.
    """
    cache_key = f"eip_{region}"
    
    if cache_key in PRICING_CACHE:
        cached_data = PRICING_CACHE[cache_key]
        if datetime.now().timestamp() - cached_data['timestamp'] < CACHE_TTL:
            return cached_data['price']
    
    try:
        pricing_client = boto3.client('pricing', region_name='us-east-1')
        location = REGION_LOCATION_MAP.get(region)
        if not location:
            raise PricingUnavailableError(f"Unknown region: {region}")
        
        response = pricing_client.get_products(
            ServiceCode='AmazonEC2',
            Filters=[
                {'Type': 'TERM_MATCH', 'Field': 'productFamily', 'Value': 'IP Address'},
                {'Type': 'TERM_MATCH', 'Field': 'location', 'Value': location},
                {'Type': 'TERM_MATCH', 'Field': 'group', 'Value': 'ElasticIP:IdleAddress'}
            ],
            MaxResults=5
        )
        
        if response['PriceList']:
            for price_item in response['PriceList']:
                price_data = json.loads(price_item)
                on_demand = price_data.get('terms', {}).get('OnDemand', {})
                if on_demand:
                    price_dimensions = list(list(on_demand.values())[0]['priceDimensions'].values())
                    for dim in price_dimensions:
                        price_per_unit = dim.get('pricePerUnit', {}).get('USD')
                        if price_per_unit and float(price_per_unit) > 0:
                            hourly_cost = float(price_per_unit)
                            PRICING_CACHE[cache_key] = {
                                'price': hourly_cost,
                                'timestamp': datetime.now().timestamp()
                            }
                            return hourly_cost
        
        raise PricingUnavailableError(f"No EIP pricing found in {region}")
        
    except PricingUnavailableError:
        raise
    except Exception as e:
        raise PricingUnavailableError(f"Failed to fetch EIP pricing in {region}: {e}")


def get_smaller_instance_type(instance_type):
    """Get one size smaller instance type, supporting more instance families"""
    # Parse instance type
    parts = instance_type.split('.')
    if len(parts) != 2:
        return None
    
    family = parts[0]
    size = parts[1]
    
    # Size progression (smallest to largest)
    size_order = ['nano', 'micro', 'small', 'medium', 'large', 'xlarge', '2xlarge', '4xlarge', 
                  '8xlarge', '9xlarge', '12xlarge', '16xlarge', '18xlarge', '24xlarge', '32xlarge', '48xlarge', 'metal']
    
    # Try to find current size in order
    try:
        current_idx = size_order.index(size)
        if current_idx > 0:
            return f"{family}.{size_order[current_idx - 1]}"
    except ValueError:
        pass
    
    # Fallback to explicit mapping for edge cases
    size_map = {
        # T2 family
        't2.2xlarge': 't2.xlarge', 't2.xlarge': 't2.large', 't2.large': 't2.medium',
        't2.medium': 't2.small', 't2.small': 't2.micro', 't2.micro': 't2.nano',
        # T3 family
        't3.2xlarge': 't3.xlarge', 't3.xlarge': 't3.large', 't3.large': 't3.medium',
        't3.medium': 't3.small', 't3.small': 't3.micro', 't3.micro': 't3.nano',
        # T3a family
        't3a.2xlarge': 't3a.xlarge', 't3a.xlarge': 't3a.large', 't3a.large': 't3a.medium',
        't3a.medium': 't3a.small', 't3a.small': 't3a.micro', 't3a.micro': 't3a.nano',
        # T4g family (Graviton)
        't4g.2xlarge': 't4g.xlarge', 't4g.xlarge': 't4g.large', 't4g.large': 't4g.medium',
        't4g.medium': 't4g.small', 't4g.small': 't4g.micro', 't4g.micro': 't4g.nano',
        # M5 family
        'm5.24xlarge': 'm5.16xlarge', 'm5.16xlarge': 'm5.12xlarge', 'm5.12xlarge': 'm5.8xlarge',
        'm5.8xlarge': 'm5.4xlarge', 'm5.4xlarge': 'm5.2xlarge', 'm5.2xlarge': 'm5.xlarge', 'm5.xlarge': 'm5.large',
        # M5a family
        'm5a.24xlarge': 'm5a.16xlarge', 'm5a.16xlarge': 'm5a.12xlarge', 'm5a.12xlarge': 'm5a.8xlarge',
        'm5a.8xlarge': 'm5a.4xlarge', 'm5a.4xlarge': 'm5a.2xlarge', 'm5a.2xlarge': 'm5a.xlarge', 'm5a.xlarge': 'm5a.large',
        # M6i family
        'm6i.32xlarge': 'm6i.24xlarge', 'm6i.24xlarge': 'm6i.16xlarge', 'm6i.16xlarge': 'm6i.12xlarge',
        'm6i.12xlarge': 'm6i.8xlarge', 'm6i.8xlarge': 'm6i.4xlarge', 'm6i.4xlarge': 'm6i.2xlarge',
        'm6i.2xlarge': 'm6i.xlarge', 'm6i.xlarge': 'm6i.large',
        # M6g family (Graviton)
        'm6g.16xlarge': 'm6g.12xlarge', 'm6g.12xlarge': 'm6g.8xlarge', 'm6g.8xlarge': 'm6g.4xlarge',
        'm6g.4xlarge': 'm6g.2xlarge', 'm6g.2xlarge': 'm6g.xlarge', 'm6g.xlarge': 'm6g.large',
        # M7i family
        'm7i.48xlarge': 'm7i.24xlarge', 'm7i.24xlarge': 'm7i.16xlarge', 'm7i.16xlarge': 'm7i.12xlarge',
        'm7i.12xlarge': 'm7i.8xlarge', 'm7i.8xlarge': 'm7i.4xlarge', 'm7i.4xlarge': 'm7i.2xlarge',
        'm7i.2xlarge': 'm7i.xlarge', 'm7i.xlarge': 'm7i.large',
        # M7a family
        'm7a.48xlarge': 'm7a.32xlarge', 'm7a.32xlarge': 'm7a.24xlarge', 'm7a.24xlarge': 'm7a.16xlarge',
        'm7a.16xlarge': 'm7a.12xlarge', 'm7a.12xlarge': 'm7a.8xlarge', 'm7a.8xlarge': 'm7a.4xlarge',
        'm7a.4xlarge': 'm7a.2xlarge', 'm7a.2xlarge': 'm7a.xlarge', 'm7a.xlarge': 'm7a.large',
        # C5 family
        'c5.24xlarge': 'c5.18xlarge', 'c5.18xlarge': 'c5.12xlarge', 'c5.12xlarge': 'c5.9xlarge',
        'c5.9xlarge': 'c5.4xlarge', 'c5.4xlarge': 'c5.2xlarge', 'c5.2xlarge': 'c5.xlarge', 'c5.xlarge': 'c5.large',
        # C5a family
        'c5a.24xlarge': 'c5a.16xlarge', 'c5a.16xlarge': 'c5a.12xlarge', 'c5a.12xlarge': 'c5a.8xlarge',
        'c5a.8xlarge': 'c5a.4xlarge', 'c5a.4xlarge': 'c5a.2xlarge', 'c5a.2xlarge': 'c5a.xlarge', 'c5a.xlarge': 'c5a.large',
        # C6i family
        'c6i.32xlarge': 'c6i.24xlarge', 'c6i.24xlarge': 'c6i.16xlarge', 'c6i.16xlarge': 'c6i.12xlarge',
        'c6i.12xlarge': 'c6i.8xlarge', 'c6i.8xlarge': 'c6i.4xlarge', 'c6i.4xlarge': 'c6i.2xlarge',
        'c6i.2xlarge': 'c6i.xlarge', 'c6i.xlarge': 'c6i.large',
        # C6g family (Graviton)
        'c6g.16xlarge': 'c6g.12xlarge', 'c6g.12xlarge': 'c6g.8xlarge', 'c6g.8xlarge': 'c6g.4xlarge',
        'c6g.4xlarge': 'c6g.2xlarge', 'c6g.2xlarge': 'c6g.xlarge', 'c6g.xlarge': 'c6g.large',
        # C7i family
        'c7i.48xlarge': 'c7i.24xlarge', 'c7i.24xlarge': 'c7i.16xlarge', 'c7i.16xlarge': 'c7i.12xlarge',
        'c7i.12xlarge': 'c7i.8xlarge', 'c7i.8xlarge': 'c7i.4xlarge', 'c7i.4xlarge': 'c7i.2xlarge',
        'c7i.2xlarge': 'c7i.xlarge', 'c7i.xlarge': 'c7i.large',
        # R5 family
        'r5.24xlarge': 'r5.16xlarge', 'r5.16xlarge': 'r5.12xlarge', 'r5.12xlarge': 'r5.8xlarge',
        'r5.8xlarge': 'r5.4xlarge', 'r5.4xlarge': 'r5.2xlarge', 'r5.2xlarge': 'r5.xlarge', 'r5.xlarge': 'r5.large',
        # R5a family
        'r5a.24xlarge': 'r5a.16xlarge', 'r5a.16xlarge': 'r5a.12xlarge', 'r5a.12xlarge': 'r5a.8xlarge',
        'r5a.8xlarge': 'r5a.4xlarge', 'r5a.4xlarge': 'r5a.2xlarge', 'r5a.2xlarge': 'r5a.xlarge', 'r5a.xlarge': 'r5a.large',
        # R6i family
        'r6i.32xlarge': 'r6i.24xlarge', 'r6i.24xlarge': 'r6i.16xlarge', 'r6i.16xlarge': 'r6i.12xlarge',
        'r6i.12xlarge': 'r6i.8xlarge', 'r6i.8xlarge': 'r6i.4xlarge', 'r6i.4xlarge': 'r6i.2xlarge',
        'r6i.2xlarge': 'r6i.xlarge', 'r6i.xlarge': 'r6i.large',
        # R6g family (Graviton)
        'r6g.16xlarge': 'r6g.12xlarge', 'r6g.12xlarge': 'r6g.8xlarge', 'r6g.8xlarge': 'r6g.4xlarge',
        'r6g.4xlarge': 'r6g.2xlarge', 'r6g.2xlarge': 'r6g.xlarge', 'r6g.xlarge': 'r6g.large',
        # I3 family (Storage optimized)
        'i3.16xlarge': 'i3.8xlarge', 'i3.8xlarge': 'i3.4xlarge', 'i3.4xlarge': 'i3.2xlarge',
        'i3.2xlarge': 'i3.xlarge', 'i3.xlarge': 'i3.large',
        # D2 family (Dense storage)
        'd2.8xlarge': 'd2.4xlarge', 'd2.4xlarge': 'd2.2xlarge', 'd2.2xlarge': 'd2.xlarge',
    }
    return size_map.get(instance_type)

def get_smaller_rds_class(db_class):
    """Get one size smaller RDS class, supporting more instance families"""
    # Parse db class
    if not db_class.startswith('db.'):
        return None
    
    parts = db_class[3:].split('.')  # Remove 'db.' prefix
    if len(parts) != 2:
        return None
    
    family = parts[0]
    size = parts[1]
    
    # Size progression (smallest to largest)
    size_order = ['micro', 'small', 'medium', 'large', 'xlarge', '2xlarge', '4xlarge', 
                  '8xlarge', '12xlarge', '16xlarge', '24xlarge', '32xlarge']
    
    # Try to find current size in order
    try:
        current_idx = size_order.index(size)
        if current_idx > 0:
            return f"db.{family}.{size_order[current_idx - 1]}"
    except ValueError:
        pass
    
    # Fallback to explicit mapping
    size_map = {
        # T3 family
        'db.t3.2xlarge': 'db.t3.xlarge', 'db.t3.xlarge': 'db.t3.large', 'db.t3.large': 'db.t3.medium',
        'db.t3.medium': 'db.t3.small', 'db.t3.small': 'db.t3.micro',
        # T4g family (Graviton)
        'db.t4g.2xlarge': 'db.t4g.xlarge', 'db.t4g.xlarge': 'db.t4g.large', 'db.t4g.large': 'db.t4g.medium',
        'db.t4g.medium': 'db.t4g.small', 'db.t4g.small': 'db.t4g.micro',
        # M5 family
        'db.m5.24xlarge': 'db.m5.16xlarge', 'db.m5.16xlarge': 'db.m5.12xlarge', 'db.m5.12xlarge': 'db.m5.8xlarge',
        'db.m5.8xlarge': 'db.m5.4xlarge', 'db.m5.4xlarge': 'db.m5.2xlarge', 'db.m5.2xlarge': 'db.m5.xlarge', 'db.m5.xlarge': 'db.m5.large',
        # M6i family
        'db.m6i.32xlarge': 'db.m6i.24xlarge', 'db.m6i.24xlarge': 'db.m6i.16xlarge', 'db.m6i.16xlarge': 'db.m6i.12xlarge',
        'db.m6i.12xlarge': 'db.m6i.8xlarge', 'db.m6i.8xlarge': 'db.m6i.4xlarge', 'db.m6i.4xlarge': 'db.m6i.2xlarge',
        'db.m6i.2xlarge': 'db.m6i.xlarge', 'db.m6i.xlarge': 'db.m6i.large',
        # M6g family (Graviton)
        'db.m6g.16xlarge': 'db.m6g.12xlarge', 'db.m6g.12xlarge': 'db.m6g.8xlarge', 'db.m6g.8xlarge': 'db.m6g.4xlarge',
        'db.m6g.4xlarge': 'db.m6g.2xlarge', 'db.m6g.2xlarge': 'db.m6g.xlarge', 'db.m6g.xlarge': 'db.m6g.large',
        # R5 family
        'db.r5.24xlarge': 'db.r5.16xlarge', 'db.r5.16xlarge': 'db.r5.12xlarge', 'db.r5.12xlarge': 'db.r5.8xlarge',
        'db.r5.8xlarge': 'db.r5.4xlarge', 'db.r5.4xlarge': 'db.r5.2xlarge', 'db.r5.2xlarge': 'db.r5.xlarge', 'db.r5.xlarge': 'db.r5.large',
        # R6i family
        'db.r6i.32xlarge': 'db.r6i.24xlarge', 'db.r6i.24xlarge': 'db.r6i.16xlarge', 'db.r6i.16xlarge': 'db.r6i.12xlarge',
        'db.r6i.12xlarge': 'db.r6i.8xlarge', 'db.r6i.8xlarge': 'db.r6i.4xlarge', 'db.r6i.4xlarge': 'db.r6i.2xlarge',
        'db.r6i.2xlarge': 'db.r6i.xlarge', 'db.r6i.xlarge': 'db.r6i.large',
        # R6g family (Graviton)
        'db.r6g.16xlarge': 'db.r6g.12xlarge', 'db.r6g.12xlarge': 'db.r6g.8xlarge', 'db.r6g.8xlarge': 'db.r6g.4xlarge',
        'db.r6g.4xlarge': 'db.r6g.2xlarge', 'db.r6g.2xlarge': 'db.r6g.xlarge', 'db.r6g.xlarge': 'db.r6g.large',
    }
    return size_map.get(db_class)


def generate_json_report(recommendations, total_savings, client_name, ri_sp_summary=None):
    """Generate a JSON report of recommendations."""
    report = {
        'client_name': client_name,
        'generated': datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC'),
        'total_monthly_savings': round(total_savings, 2),
        'total_annual_savings': round(total_savings * 12, 2),
        'recommendations': {},
        'ri_sp_coverage': ri_sp_summary
    }
    
    for service, recs in recommendations.items():
        if isinstance(recs, list):
            report['recommendations'][service] = recs
            report['recommendations'][f'{service}_count'] = len(recs)
            report['recommendations'][f'{service}_savings'] = round(sum(r.get('monthly_savings', 0) for r in recs), 2)
    
    return json.dumps(report, indent=2, default=str)


def generate_csv_report(recommendations, total_savings, client_name):
    """Generate a CSV report of recommendations."""
    output = StringIO()
    writer = csv.writer(output)
    
    # Header
    writer.writerow(['Service', 'Resource ID', 'Tags', 'Region', 'Issue/Reason',
                      'Recommendation', 'Current Cost', 'Recommended Cost',
                      'Monthly Savings', 'Annual Savings', 'Confidence'])
    
    service_configs = {
        'ec2': lambda r: [r['instance_id'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                          r.get('reason', ''), f"Resize to {r['recommended_type']}",
                          f"${r['current_cost']:.2f}", f"${r['recommended_cost']:.2f}",
                          f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                          r['confidence']],
        'stopped_ec2': lambda r: [r['instance_id'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                                   r.get('reason', ''), r['recommendation'],
                                   '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                                   r['confidence']],
        'ebs': lambda r: [r['volume_id'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                          r['issue'], r['recommendation'],
                          '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                          r['confidence']],
        'rds': lambda r: [r['db_id'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                          r.get('reason', ''), f"Resize to {r['recommended_class']}",
                          f"${r['current_cost']:.2f}", f"${r['recommended_cost']:.2f}",
                          f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                          r['confidence']],
        'lambda': lambda r: [r['function_name'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                             f"Memory: {r['current_memory']}MB", f"Reduce to {r['recommended_memory']}MB",
                             f"${r['current_cost']:.2f}", f"${r['recommended_cost']:.2f}",
                             f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                             r['confidence']],
        'eip': lambda r: [r['ip_address'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                          r['status'], r['recommendation'],
                          '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                          r['confidence']],
        'natgateway': lambda r: [r['nat_gateway_id'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                                  r.get('reason', ''), r['recommendation'],
                                  '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                                  r['confidence']],
        's3': lambda r: [r['bucket_name'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                         r['issues'], r['recommendation'],
                         '', '', '$0.00', '$0.00', r['confidence']],
        'dynamodb': lambda r: [r['table_name'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                               r.get('reason', ''), r['recommendation'],
                               f"${r['current_cost']:.2f}", f"${r['recommended_cost']:.2f}",
                               f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                               r['confidence']],
        'ebs_snapshot': lambda r: [r['snapshot_id'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                                   r.get('issue', ''), r['recommendation'],
                                   '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                                   r['confidence']],
        'elb': lambda r: [r['load_balancer_name'], format_tags_str(r.get('tags', {})), r.get('region', ''),
                          r.get('reason', ''), r['recommendation'],
                          '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                          r['confidence']],
        'savings_plan': lambda r: [r.get('type', 'Savings Plan'), '', 'global',
                                   r.get('reason', ''), r['recommendation'],
                                   '', '', f"${r['monthly_savings']:.2f}", f"${r['monthly_savings'] * 12:.2f}",
                                   r['confidence']],
    }
    
    for service, recs in recommendations.items():
        if isinstance(recs, list) and service in service_configs:
            for rec in recs:
                try:
                    row_data = service_configs[service](rec)
                    writer.writerow([service.upper()] + row_data)
                except KeyError as e:
                    # Write error row with correct number of columns (10 data columns)
                    error_row = [f'Error: missing field {e}'] + [''] * 9
                    writer.writerow([service.upper()] + error_row)
    
    # Summary row
    writer.writerow([])
    writer.writerow(['TOTAL', '', '', '', '', '', '',
                     f'${total_savings:.2f}', f'${total_savings * 12:.2f}', ''])
    
    return output.getvalue()


def _rec_resource_id(rec):
    """Best-effort human resource identifier for a recommendation dict."""
    for key in ('instance_id', 'volume_id', 'db_id', 'function_name', 'snapshot_id',
                'load_balancer_name', 'nat_gateway_id', 'table_name', 'ip_address',
                'bucket_name', 'type'):
        if rec.get(key):
            return rec[key]
    return 'resource'


def generate_html_report(result, client_name):
    """Generate a self-contained interactive HTML dashboard (string)."""
    try:
        import dashboard_assets
    except ImportError:  # pragma: no cover
        from . import dashboard_assets  # type: ignore
    data = scan_result_summary(result)
    data['clientName'] = client_name
    return dashboard_assets.build_standalone_html(json.dumps(data, default=str), client_name)


def generate_xlsx_report(result, client_name):
    """Generate a multi-sheet Excel workbook (bytes): Summary + per-service + RI/SP."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("XLSX export requires the 'openpyxl' package.") from e

    recommendations = result['recommendations']
    header_fill = PatternFill('solid', fgColor='23649C')
    header_font = Font(bold=True, color='FFFFFF')

    def _style_header(ws, row_idx=1):
        for cell in ws[row_idx]:
            cell.font = header_font
            cell.fill = header_fill

    def _safe_title(name):
        for ch in '[]:*?/\\':
            name = name.replace(ch, '-')
        return name[:31]

    wb = Workbook()
    ws = wb.active
    ws.title = 'Summary'
    ws.append(['CostOptimizer360 - AWS Cost Optimization Report'])
    ws['A1'].font = Font(bold=True, size=14, color='1A4D78')
    ws.append(['Client', client_name])
    ws.append(['Generated', result.get('generated_at', '')])
    ws.append(['Monthly Savings (waste elimination)', result['total_savings']])
    ws.append(['Annual Savings (waste elimination)', round(result['total_savings'] * 12, 2)])
    ws.append(['Savings Plans opportunity (monthly)', result.get('commitment_savings', 0)])
    ws.append(['Quick Wins', result['metrics'].get('quick_wins', 0)])
    ws.append(['High Priority', result['metrics'].get('high_priority', 0)])
    forecast = result.get('forecast') or {}
    if forecast.get('forecast_month') is not None:
        ws.append(['Forecast spend (this month)', forecast['forecast_month']])
    ws.append([])
    ws.append(['Service', 'Recommendations', 'Monthly Savings', 'Annual Savings'])
    _style_header(ws, ws.max_row)
    for svc, recs in recommendations.items():
        if isinstance(recs, list) and recs:
            svc_savings = sum(r.get('monthly_savings', 0) for r in recs)
            ws.append([SERVICE_LABELS.get(svc, svc), len(recs),
                       round(svc_savings, 2), round(svc_savings * 12, 2)])

    # Per-service detail sheets
    for svc, recs in recommendations.items():
        if not (isinstance(recs, list) and recs):
            continue
        wsx = wb.create_sheet(title=_safe_title(SERVICE_LABELS.get(svc, svc)))
        wsx.append(['Resource', 'Region', 'Reason/Issue', 'Recommendation',
                    'Monthly Savings', 'Annual Savings', 'Confidence',
                    'Priority', 'Effort', 'Risk', 'Savings Basis'])
        _style_header(wsx)
        for r in recs:
            wsx.append([
                _rec_resource_id(r), r.get('region', ''),
                r.get('reason') or r.get('issue') or r.get('issues') or '',
                r.get('recommendation', ''),
                r.get('monthly_savings', 0), r.get('annual_savings', 0),
                r.get('confidence', ''), r.get('priority', ''),
                r.get('effort', ''), r.get('risk', ''), r.get('savings_basis', ''),
            ])

    # RI/SP coverage sheet
    ri_sp = result.get('ri_sp_summary') or {}
    if ri_sp:
        wsr = wb.create_sheet(title='RI-SP Coverage')
        wsr.append(['Metric', 'Value'])
        _style_header(wsr)
        wsr.append(['Running instances', ri_sp.get('total_running_instances', 0)])
        wsr.append(['RI coverage %', ri_sp.get('ri_coverage_pct', 0)])
        wsr.append(['RI coverage basis', ri_sp.get('ri_coverage_basis', 'n/a')])
        wsr.append(['Savings Plans coverage %', ri_sp.get('sp_coverage_pct', 0)])
        wsr.append(['Savings Plans utilization %', ri_sp.get('sp_utilization_pct', 'n/a')])

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.read()


# Services scanned once per region vs global (S3/commitments/forecast handled separately)
PER_REGION_SCANNERS = ['ec2', 'stopped_ec2', 'ebs', 'ebs_snapshot', 'rds',
                       'lambda', 'eip', 'natgateway', 'dynamodb', 'elb']

SERVICE_LABELS = {
    'ec2': 'EC2 Instances', 'stopped_ec2': 'Stopped EC2 Instances',
    'ebs': 'EBS Volumes', 'ebs_snapshot': 'EBS Snapshots', 'rds': 'RDS Databases',
    'lambda': 'Lambda Functions', 'eip': 'Elastic IPs / Public IPv4',
    'natgateway': 'NAT Gateways', 'dynamodb': 'DynamoDB Tables',
    'elb': 'Load Balancers', 's3': 'S3 Buckets', 'savings_plan': 'Savings Plans',
}


def resolve_regions(body):
    """Resolve the list of regions to scan from the request body."""
    regions_input = body.get('regions', body.get('region', 'us-east-1'))
    if isinstance(regions_input, str):
        if regions_input == 'all':
            temp_session = create_session(body, 'us-east-1')
            ec2_client = temp_session.client('ec2')
            return [r['RegionName'] for r in ec2_client.describe_regions()['Regions']]
        return [regions_input]
    return regions_input


def run_full_scan(body, progress_cb=None):
    """Run every requested scan and return a normalized result dict.

    Shared by both the Lambda handler and the local Flask server so the scan
    logic lives in exactly one place. `progress_cb(step, total_steps, label)` is
    called before each unit of work for progress reporting.

    Savings are split into two honest buckets so we never conflate them:
      * total_savings       - waste elimination (rightsizing, idle, storage...)
      * commitment_savings  - rate optimization (Savings Plans purchase)
    """
    services = body.get('services', ['ec2', 'ebs', 'rds', 'lambda', 'eip'])
    regions = resolve_regions(body)

    scanners = {
        'ec2': scan_ec2_instances,
        'stopped_ec2': scan_stopped_ec2_instances,
        'ebs': scan_ebs_volumes,
        'ebs_snapshot': scan_ebs_snapshots,
        'rds': scan_rds_instances,
        'lambda': scan_lambda_functions,
        'eip': scan_elastic_ips,
        'natgateway': scan_nat_gateways,
        'dynamodb': scan_dynamodb_tables,
        'elb': scan_load_balancers,
    }

    per_region_selected = [s for s in services if s in scanners]
    total_steps = max(1, len(regions) * len(per_region_selected)
                      + (1 if 's3' in services else 0)
                      + (1 if 'ec2' in services else 0)  # RI/SP coverage
                      + (1 if 'commitments' in services else 0))
    step = 0

    recommendations = {}
    total_savings = 0.0

    for region in regions:
        session = create_session(body, region)
        for svc in per_region_selected:
            step += 1
            if progress_cb:
                region_label = f" ({region})" if len(regions) > 1 else ""
                progress_cb(step, total_steps, f"{SERVICE_LABELS.get(svc, svc)}{region_label}")
            try:
                recs = scanners[svc](session)
            except Exception as e:
                print(f"Scanner {svc} failed in {region}: {e}")
                recs = []
            for r in recs:
                r['region'] = region
            recommendations.setdefault(svc, []).extend(recs)
            total_savings += sum(r.get('monthly_savings', 0) for r in recs)

    # S3 is global - scan once.
    if 's3' in services:
        step += 1
        if progress_cb:
            progress_cb(step, total_steps, SERVICE_LABELS['s3'])
        try:
            recommendations['s3'] = scan_s3_buckets(create_session(body, 'us-east-1'))
            total_savings += sum(r.get('monthly_savings', 0) for r in recommendations['s3'])
        except Exception as e:
            print(f"S3 scan failed: {e}")

    # RI/SP coverage - from first region.
    ri_sp_summary = None
    if 'ec2' in services and regions:
        step += 1
        if progress_cb:
            progress_cb(step, total_steps, 'RI / Savings Plans Coverage')
        try:
            ri_sp_summary = scan_ri_sp_coverage(create_session(body, regions[0]))
        except Exception as e:
            print(f"RI/SP coverage failed: {e}")

    # Commitment (rate optimization) recommendations - kept in a separate bucket.
    commitment_savings = 0.0
    if 'commitments' in services:
        step += 1
        if progress_cb:
            progress_cb(step, total_steps, 'Savings Plans Purchase Analysis')
        try:
            sp_recs = scan_savings_plans_purchase(create_session(body, 'us-east-1'))
            if sp_recs:
                for r in sp_recs:
                    r['region'] = 'global'
                recommendations['savings_plan'] = sp_recs
                commitment_savings += sum(r.get('monthly_savings', 0) for r in sp_recs)
        except Exception as e:
            print(f"Savings Plans purchase analysis failed: {e}")

    # Cost forecast + month-to-date spend for the dashboard/report.
    forecast = None
    if regions:
        try:
            forecast = get_cost_forecast_and_spend(create_session(body, 'us-east-1'))
        except Exception as e:
            print(f"Forecast lookup failed: {e}")

    # Attach effort/risk/priority/annualized savings/remediation to every rec.
    metrics = enrichment.enrich_recommendations(recommendations)

    return {
        'recommendations': recommendations,
        'total_savings': round(total_savings, 2),
        'commitment_savings': round(commitment_savings, 2),
        'ri_sp_summary': ri_sp_summary,
        'forecast': forecast,
        'metrics': metrics,
        'regions': regions,
        'services': services,
        'generated_at': datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC'),
    }


def make_report(result, client_name, export_format):
    """Generate the requested report artifact. Returns (bytes, filename)."""
    recommendations = result['recommendations']
    total = result['total_savings']
    ri_sp = result['ri_sp_summary']
    date = datetime.now(timezone.utc).strftime('%Y%m%d')
    base = f"{client_name.replace(' ', '-')}-CostOptimizer360-{date}"

    if export_format == 'json':
        return generate_json_report(recommendations, total, client_name, ri_sp).encode('utf-8'), base + '.json'
    if export_format == 'csv':
        return generate_csv_report(recommendations, total, client_name).encode('utf-8'), base + '.csv'
    if export_format == 'html':
        return generate_html_report(result, client_name).encode('utf-8'), base + '.html'
    if export_format == 'xlsx':
        return generate_xlsx_report(result, client_name), base + '.xlsx'

    # Default: Word document
    doc = generate_word_report(recommendations, total, client_name, ri_sp, result)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read(), base + '.docx'


def scan_result_summary(result):
    """Build the JSON-serializable summary returned to the frontend (feeds the dashboard)."""
    recommendations = result['recommendations']
    return {
        'totalMonthlySavings': result['total_savings'],
        'totalAnnualSavings': round(result['total_savings'] * 12, 2),
        'commitmentMonthlySavings': result.get('commitment_savings', 0.0),
        'commitmentAnnualSavings': round(result.get('commitment_savings', 0.0) * 12, 2),
        'recommendationCounts': {k: len(v) for k, v in recommendations.items() if isinstance(v, list)},
        'riSpCoverage': result['ri_sp_summary'],
        'forecast': result.get('forecast'),
        'quickWins': result['metrics'].get('quick_wins', 0),
        'highPriority': result['metrics'].get('high_priority', 0),
        'regions': result.get('regions', []),
        'generatedAt': result.get('generated_at'),
        'recommendations': recommendations,
    }


def create_session(body, region):
    """Create a boto3 session for the given region.

    Cross-account access uses STS AssumeRole with an ExternalId (confused-deputy
    protection). The target-account trust policy requires this ExternalId, so it
    MUST be passed or the assume-role call fails with AccessDenied. The value is
    configurable per tenant; it defaults to the documented shared value only for
    backward compatibility.
    """
    if 'roleArn' in body:
        sts = boto3.client('sts')
        external_id = body.get('externalId') or os.environ.get('EXTERNAL_ID', 'CostOptimizer360')
        assume_kwargs = {
            'RoleArn': body['roleArn'],
            'RoleSessionName': 'CostOptimizer360Session',
            'ExternalId': external_id,
            'DurationSeconds': 3600,
        }
        assumed_role = sts.assume_role(**assume_kwargs)
        return boto3.Session(
            aws_access_key_id=assumed_role['Credentials']['AccessKeyId'],
            aws_secret_access_key=assumed_role['Credentials']['SecretAccessKey'],
            aws_session_token=assumed_role['Credentials']['SessionToken'],
            region_name=region
        )
    else:
        session_kwargs = {
            'aws_access_key_id': body['accessKeyId'],
            'aws_secret_access_key': body['secretAccessKey'],
            'region_name': region
        }
        if body.get('sessionToken'):
            session_kwargs['aws_session_token'] = body['sessionToken']
        return boto3.Session(**session_kwargs)
